import streamlit as st
import pandas as pd
from io import BytesIO
import pdfplumber
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
import re
import datetime
import json
from rapidfuzz import process, fuzz
import os
from reportlab.platypus import Image
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.pagesizes import letter
import html

# ----------------------------
# Load JSON reference files
# ----------------------------
with open("material_grades.json", "r", encoding="utf-8") as f:
    raw_data = json.load(f)

def flatten_material_grades(data):
    all_entries = []
    if isinstance(data, dict):
        if any(k.lower() in ["material no.", "grade", "combined_grade", "finish grade", "grade specification", "usa", "germany din", "gb bs", "japan", "identification/marking"] for k in data.keys()):
            all_entries.append(data)
        for v in data.values():
            all_entries.extend(flatten_material_grades(v))
    elif isinstance(data, list):
        for item in data:
            all_entries.extend(flatten_material_grades(item))
    return all_entries

material_data = flatten_material_grades(raw_data)

# Extract all distinct grade-like values
material_grades = set()
for entry in material_data:
    for key, val in entry.items():
        if isinstance(val, str) and val.strip():
            material_grades.add(val.strip().upper())
material_grades = list(material_grades)

with open("DIN_standards.json", "r", encoding="utf-8") as f:
    din_data = json.load(f)
    din_standards = list({x["Grade specification"] for x in din_data if "Grade specification" in x})


# ----------------------------
# Quotation schema
# ----------------------------
quotation_columns = [
    "Sr. No.", "Item Code", "Material Description", "Dimension\nStandard",
    "Material\nGrade", "Finish", "Qty/MOQ", "Rate"
]

# ----------------------------
# Helpers
# ----------------------------
def _norm(s):
    if s is None: return ""
    s = re.sub(r'\s+', ' ', str(s))
    s = re.sub(r'[^\x20-\x7E]', '', s)
    return s.strip()

def validate_multi(val, valid_list):
    if not val: return ""
    parts = [p.strip() for p in str(val).split(",")]
    validated = []
    for p in parts:
        if not p: continue
        match = process.extractOne(_norm(p), valid_list, scorer=fuzz.ratio)
        if match:
            m, score, _ = match
            if score >= 80:
                validated.append(m)
    return ", ".join(validated)

item_code_pattern = re.compile(r"\b[A-Z0-9]+(?:[-/_\.]?[A-Z0-9]+){0,3}\b", re.I)

def clean_line(text: str) -> str:
    text = re.sub(r"^\s*\d+(\s|$)", " ", str(text))
    text = re.sub(r"\s+", " ", text).strip()
    return text

def is_header_line(text: str) -> bool:
    header_keywords = [
        "material", "description", "sr", "qty", "moq", "rate",
        "item", "code", "finish", "grade", "standard", "dimension"
    ]
    text_lower = text.lower()
    return any(k in text_lower for k in header_keywords)

def clean_material_description(desc: str) -> str:
    desc = re.sub(r"\s+", " ", desc)
    desc = re.sub(r"\s*x\s*", " x ", desc, flags=re.I)
    return desc.strip()

def find_grade_reference_details(desc_text, material_data):
    if not desc_text:
        return []
    desc_upper = desc_text.upper()
    matched_refs = []
    for entry in material_data:
        for key, val in entry.items():
            if isinstance(val, str) and val.strip():
                if val.upper() in desc_upper:
                    matched_refs.append(entry)
                    break
    return matched_refs

def parse_line(text: str):
    text = clean_line(text)
    if not text or is_header_line(text):
        return None
    tokens = text.split()
    item_code = ""
    for t in tokens:
        if len(t) >= 3 and re.fullmatch(item_code_pattern, t):
            if not re.fullmatch(r"(WASHER|NUT|BOLT|CAPSCREW|SCREW|STUD|PIN|ROD|HEAD|HEX)", t, re.I):
                item_code = t
                break
    desc_part = text.replace(item_code, "", 1).strip() if item_code else text
    nums = re.findall(r"\b\d+(?:\.\d+)?\b", desc_part)
    qty = nums[-1] if nums else ""
    desc_clean = re.sub(r"(\b\d+(?:\.\d+)?\b\s*)$", "", desc_part).strip()
    desc_clean = clean_material_description(desc_clean)
    return {
        "Item Code": item_code,
        "Material Description": desc_clean,
        "Dimension\nStandard": "",
        "Material\nGrade": "",
        "Finish": "",
        "Qty/MOQ": qty,
        "Rate": ""  # Always blank
    }

def parse_material_description(row):
    desc = row.get("Material Description", "") or ""
    desc_clean = clean_material_description(desc).upper()
    found_grade, found_finish, found_dim = set(), set(), set()
    DIM_REGEX = r"(M\d+(\s*X\s*\d+P)?\s*X\s*\d+L)|FULL THREAD|DIN\s*\d+|UNC|UNF|\d+/\d+INCH|\d+IN"
    parts = [p.strip() for p in desc_clean.split(",") if p.strip()]
    remaining_desc = []

    for part in parts:
        # Finish detection
        for f in din_standards:
            if fuzz.partial_ratio(f.upper(), part) >= 85:
                found_finish.add(f)

        # Grade detection (fuzzy match against all material_grades)
        for g in material_grades:
            if fuzz.partial_ratio(g, part) >= 90:
                found_grade.add(g)

        # Dimension detection
        dims = re.findall(DIM_REGEX, part)
        for d in dims:
            found_dim.add(d[0] if isinstance(d, tuple) else d)

        cleaned_part = part
        for term in list(found_finish) + list(found_grade) + list(found_dim):
            cleaned_part = re.sub(re.escape(term), "", cleaned_part)
        cleaned_part = re.sub(r'\s+', ' ', cleaned_part).strip()
        if cleaned_part:
            remaining_desc.append(cleaned_part)

    # Fill row fields
    row["Material Description"] = " | ".join(remaining_desc)
    row["Finish"] = ", ".join(found_finish)
    row["Material\nGrade"] = ", ".join(found_grade)
    row["Dimension\nStandard"] = ", ".join(found_dim)

    # --- Reference lookup ---
    matched_refs = find_grade_reference_details(desc_clean, material_data)
    if matched_refs:
        row["_grade_details"] = matched_refs
        # Fill grade column if still empty
        if not row["Material\nGrade"].strip():
            all_vals = {val for ref in matched_refs for val in ref.values() if val and isinstance(val, str)}
            row["Material\nGrade"] = ", ".join(sorted(all_vals))

    return row

SAVE_FILE = "saved_values.json"

# ---------------------------
# LOAD SAVED VALUES
# ---------------------------
def load_saved_values():
    if os.path.exists(SAVE_FILE):
        try:
            with open(SAVE_FILE, "r") as f:
                return json.load(f)
        except:
            return {}        # empty or corrupt file
    return {}

# ---------------------------
# SAVE VALUES
# ---------------------------
def save_values(data):
    with open(SAVE_FILE, "w") as f:
        json.dump(data, f, indent=4)


# -------------------------------------------------------------
# SINGLE ROW INPUT WITH SMART SUGGESTIONS + AUTO-SAVE
# -------------------------------------------------------------
def smart_input(label, key_name, saved_dict):
    # Load existing saved values
    options = saved_dict.get(key_name, [])

    # --- Only ONE visible row ---
    typed = st.text_input(label, key=f"{key_name}_box")
    st.session_state["header_data"][key_name] = typed

    # Show suggestions dynamically
    matches = [o for o in options if typed.lower() in o.lower()] if typed else []

    if matches:
        selected = st.selectbox(
            f"Select existing {label}",
            matches,
            key=f"{key_name}_suggest"
        )
        if selected:
            return selected

    # Auto-save new value
    if typed.strip() and typed not in options:
        options.append(typed)
        saved_dict[key_name] = options
        save_values(saved_dict)

    return typed

# ---- GLOBAL OPTIONS (Always available + Save new values) ----

if "delivery_options" not in st.session_state:
    st.session_state.delivery_options = set([
        "EX-MUMBAI,CUSTOMER TRANSPORT",
        "EX WORKS",
        "DELIVERY INCLUSIVE",
        "GO DOWN DELIVERY"
    ])

if "payment_options" not in st.session_state:
    st.session_state.payment_options = set([
        "IMMEDIATE",
        "30% ADVANCE BALANCE AGAINST PI",
        "30 DAYS CREDIT",
        "45 DAYS CREDIT",
        "100% ADVANCE AGAINST PI"
    ])

if "validity_options" not in st.session_state:
    st.session_state.validity_options = set([
        "1 DAY",
        "1 WEEK",
        "1 MONTH"
    ])


# --------------------------------------------
# Extraction Functions: PDF / Excel / JSON
def extract_pdf(file):
    parsed_rows = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            if not tables:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    parsed = parse_line(line.strip())
                    if parsed:
                        parsed_rows.append(parsed)
                continue

            for table in tables:
                rows = [[str(c).strip() if c else "" for c in row] for row in table]
                header_idx = None
                for i, row in enumerate(rows):
                    join = " ".join(row).lower().replace(" ", "").replace("_", "")
                    if any(k in join for k in ["item", "desc", "qty", "finish", "grade", "standard", "dimension"]):
                        header_idx = i
                        break
                if header_idx is None:
                    for row in rows:
                        text = " ".join(row)
                        parsed = parse_line(text)
                        if parsed:
                            parsed_rows.append(parsed)
                    continue

                headers = rows[header_idx]
                headers_norm = [re.sub(r"[\s_\.]+", "", h.lower()) for h in headers]
                for r in rows[header_idx+1:]:
                    if not any(r): continue
                    row_dict = {k: "" for k in quotation_columns[1:]}
                    for ci, cell in enumerate(r):
                        if ci >= len(headers_norm): continue
                        val = cell.strip()
                        h = headers_norm[ci]
                        if not val: continue
                        if h in ["srno", "sr", "no"]: continue
                        elif any(k in h for k in ["itemcode", "itemnumber", "code","Item"]):
                            row_dict["Item Code"] = val
                        elif any(k in h for k in ["description", "materialdesc", "itemdesc"]):
                            row_dict["Material Description"] += " " + val
                        elif any(k in h for k in ["dimension", "standard", "std"]):
                            row_dict["Dimension\nStandard"] = val
                        elif any(k in h for k in ["grade", "materialgrade"]):
                            row_dict["Material\nGrade"] = val
                        elif "finish" in h:
                            row_dict["Finish"] = val
                        elif any(k in h for k in ["qty", "quantity", "moq", "set"]):
                            row_dict["Qty/MOQ"] = val
                        # FIXED - Do NOT map any rate/price/amount values
                    row_dict["Rate"] = ""  # always blank
                    row_dict["Material Description"] = clean_material_description(row_dict["Material Description"])
                    row_dict = parse_material_description(row_dict)
                    parsed_rows.append(row_dict)

    df = pd.DataFrame(parsed_rows).fillna("")
    if not df.empty:
        df.insert(0, "Sr. No.", range(1, len(df)+1))
        df = df.reindex(columns=quotation_columns)
    return df

# --------------------------------------------
# Excel Extraction
def extract_excel(file):
    df_raw = pd.read_excel(file, dtype=str, header=None).fillna("")
    header_idx = None

    # 🔍 Identify header row
    for i, row in df_raw.iterrows():
        join = " ".join(str(x).lower().replace("_", "").replace(" ", "") for x in row)
        if any(k in join for k in ["itemcode", "itemnumber", "itemno", "item", "code"]):
            header_idx = i
            break

    # 🔍 Normalize header names
    if header_idx is not None:
        df = pd.read_excel(file, dtype=str, header=header_idx).fillna("")
        cols_clean = [re.sub(r'[\s_\.]+', '', str(c).lower()) for c in df.columns]

        # ✅ Detect item code column (more flexible)
        has_code = any(
            any(k in c for k in ["itemcode", "itemnumber", "itemno", "item", "code"])
            for c in cols_clean
        )

        # ✅ Detect quantity column
        qty_col = None
        for col, norm in zip(df.columns, cols_clean):
            if any(k in norm for k in ["qty", "quantity", "moq", "set"]):
                qty_col = col
                break
    else:
        df = df_raw
        cols_clean = []
        has_code = False
        qty_col = None

    rows = []
    for _, row in df.iterrows():
        text = " ".join([str(x).strip() for x in row if str(x).strip()])
        parsed = parse_line(text)

        if parsed:
            # ✅ Extract item code properly even if header is 'Item'
            if has_code:
                for col, norm in zip(df.columns, cols_clean):
                    if any(k in norm for k in ["itemcode", "itemnumber", "itemno", "item", "code"]):
                        parsed["Item Code"] = str(row[col]).strip()
                        break
            else:
                parsed["Item Code"] = ""

            # ✅ Extract Qty properly
            if qty_col:
                parsed["Qty/MOQ"] = str(row[qty_col]).strip()
            else:
                parsed["Qty/MOQ"] = parsed.get("Qty/MOQ", "")

            parsed["Rate"] = ""
            rows.append(parsed)

    final_df = pd.DataFrame(rows)
    if not final_df.empty:
        final_df.insert(0, "Sr. No.", range(1, len(final_df) + 1))
        final_df = final_df.reindex(columns=quotation_columns)

    return final_df

# --------------------------------------------
# JSON Extraction
def extract_json(file):
    data = json.load(file)
    text_data = json.dumps(data, indent=2)
    lines = [l.strip() for l in text_data.splitlines() if l.strip()]
    rows = []
    for idx, line in enumerate(lines, start=1):
        parsed = parse_line(line)
        if parsed:
            parsed["Sr. No."] = idx
            parsed["Rate"] = ""  # FIXED
            rows.append(parsed)
    df = pd.DataFrame(rows).fillna("")
    df = df.reindex(columns=quotation_columns)
    return df

def generate_pdf_dynamic(quotation):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=20, leftMargin=20,
        topMargin=30, bottomMargin=30
    )

    styles = getSampleStyleSheet()
    styleN = styles['Normal']
    styleN.allowHTML = True             
    styleB = ParagraphStyle('Bold', parent=styleN, fontName='Helvetica-Bold')
    styleB.allowHTML = True             
    styleWrap = ParagraphStyle('Wrap', parent=styleN, wordWrap='CJK')
    styleWrap.allowHTML = True          


    elements = []
    # ---------------------------------------
    # COMPANY LOGO AT TOP
    # ---------------------------------------
    try:
       logo_path = "logo.jpeg"
       logo = Image(logo_path, width=520, height=100)  # adjust size as required
       elements.append(logo)
       elements.append(Spacer(1, 12))
    except Exception as e:
       print("Logo load error:", e)
    
    title_style = ParagraphStyle(
       'Title',
        parent=styles['Heading1'],
        alignment=1,      
        fontName='Helvetica-Bold',
        fontSize=18,
        spaceAfter=12
    )
    elements.append(Paragraph("QUOTATION", title_style))
    elements.append(Spacer(1, 18))

    h = quotation["header"]

# --------------------------------------------------------
# PREMIUM OPTION 5 — EXECUTIVE TWIN PANEL HEADER
# --------------------------------------------------------

# LEFT PANEL (REF)
    left_panel = Table(
      [
        [Paragraph("<b>REF NO.</b>", styleN)],
        [Paragraph(h.get("quotation_no", "NA"), styleB)]
      ],
      colWidths=[260]
    )
    left_panel.setStyle(TableStyle([
    ('BOX', (0,0), (-1,-1), 1, colors.black),
    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ('TOPPADDING', (0,0), (-1,-1), 6),
    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))

# RIGHT PANEL (DATE)
    right_panel = Table(
       [
        [Paragraph("<b>DATE</b>", styleN)],
        [Paragraph(h.get("date", "NA").replace("DATE : ", ""), styleB)]
       ],
        colWidths=[260]
    )
    right_panel.setStyle(TableStyle([
    ('BOX', (0,0), (-1,-1), 1, colors.black),
    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ('TOPPADDING', (0,0), (-1,-1), 6),
    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
     ]))

# COMBINED PANEL TABLE
    dual_panel = Table(
      [[left_panel, right_panel]],
       colWidths=[270, 270]
    )
    dual_panel.setStyle(TableStyle([
    ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ('BOTTOMPADDING', (0,0), (-1,-1), 12),
    ]))

    elements.append(dual_panel)
    elements.append(Spacer(1, 14))


# --------------------------------------------------------
# PREMIUM CLIENT INFORMATION BLOCK (Stylish + Bold)
# --------------------------------------------------------
    client_values = [
        h.get("client_name", "NA"),
        h.get("client_address", "NA"),
        h.get("mobile_no", "NA"),
        h.get("pan_no", "NA"),
        h.get("pincode", "NA"),
        h.get("state", "NA"),
        h.get("email_id", "NA"),
        h.get("gst_no", "NA")
    ]

# Ensure it becomes exactly 9 items
    while len(client_values) < 9:
       client_values.append("")

    client_grid = [
       [
        Paragraph(str(client_values[i]), styleN),
        Paragraph(str(client_values[i+1]), styleN),
        Paragraph(str(client_values[i+2]), styleN)
      ]
      for i in range(0, 9, 3)
    ]

    client_table = Table(
       client_grid,
       colWidths=[180,180,180]   # EXACT same width distribution as REF/DATE combined table
    )

    client_table.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 1, colors.black),      
        ('INNERGRID', (0,0), (-1,-1), 0, colors.white), 
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),

        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))

    elements.append(client_table)
    elements.append(Spacer(1, 14))



# --------------------------------------------------------
# INTRO PARAGRAPH
# --------------------------------------------------------
    elements.append(Paragraph(quotation.get("intro", ""), styleN))
    elements.append(Spacer(1, 12))
    # ----------------------------
    # TABLE
    # ----------------------------
    data = [quotation["items_table"]["headers"]] + quotation["items_table"]["rows"]
    col_widths = [40, 70, 180, 70, 70, 60, 50, 50]

    # SAFELY CONVERT any cell into printable text
    def safe_text(x):
        if isinstance(x, list):
            return ", ".join(str(i) for i in x)
        return str(x)
    def _normalize_cell_for_pdf(cell):
        """Ensure cell is a plain string, unescaped, and uses <font> for color."""
        s = safe_text(cell)
        # unescape HTML entities (if stored as &lt;font&gt; etc.)
        s = html.unescape(s)
        # convert any span color markers (if present) to <font> which ReportLab accepts
        s = s.replace("<span style='color:red'>", "<font color='red'>").replace("<span style=\"color:red\">", "<font color='red'>")
        s = s.replace("</span>", "</font>")
        return s

    wrapped_data = []
    for row in data:
        wrapped_row = []
        for i, cell in enumerate(row):
            text = _normalize_cell_for_pdf(cell)
            wrapped_row.append(Paragraph(text, styleWrap if i == 2 else styleN))
        wrapped_data.append(wrapped_row)

    table = Table(wrapped_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (6, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 14))
    # -----------------------------------
    # ADDITIONAL NOTE (optional field)
    # -----------------------------------
    
    additional_note = quotation.get("additional_note", "").strip()

    if additional_note:
        elements.append(Paragraph("<b>ADDITIONAL NOTE</b>", styleB))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph(f"<font color='red'>{additional_note}</font>", styleN))
        elements.append(Spacer(1, 12))

# ----------------------------------------------------
# COMPACT TOTALS BLOCK (fits neatly after 5th column)
# ----------------------------------------------------
    totals = quotation.get("totals", {})

    if totals.get("show", True):

        subtotal = totals.get("subtotal")
        discount_percent = totals.get("discount_percent")
        discount_amount = totals.get("discount_amount")
        grand_total = totals.get("grand_total")

    # Format numbers with non-breaking commas to prevent wrapping
        def format_amount(amount):
            return f"{amount:,.2f}".replace(",", "\u00A0,")  # non-breaking comma

        rows = []

    # Subtotal
        if subtotal is not None:
            rows.append([
                
                Paragraph("<b>Subtotal:</b>", styleN),
                Paragraph(f"₹ {format_amount(subtotal)}", styleB)
            ])

    # Discount
        if discount_percent is not None and discount_amount is not None:
            rows.append([
                
                Paragraph(f"<b>Discount ({discount_percent}%):</b>", styleN),
                Paragraph(f"- ₹ {format_amount(discount_amount)}", styleB)
            ])

    # Grand Total
        if grand_total is not None:
            rows.append([
                
                Paragraph("<b>Grand Total:</b>", styleN),
                Paragraph(f"₹ {format_amount(grand_total)}", styleB)
            ])
        totals_table = Table(rows, colWidths=[None, None], hAlign='RIGHT')
        totals_table.setStyle(TableStyle([
        ('ALIGN', (5, 0), (-1, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))

        elements.append(totals_table)
        elements.append(Spacer(1, 10))


# ----------------------------
# NOTES SECTION (plain text, no table)
# ----------------------------
    notes = quotation.get("notes", {})
    main_note = notes.get("main_note", "")
    sub_notes = notes.get("sub_notes", [])

    if main_note:
        elements.append(Paragraph(f"<b>{main_note}</b>", styleN))
        elements.append(Spacer(1, 6))

    for n in sub_notes:
        if "highlighted changes" in n.lower():
            elements.append(Paragraph(f"<font color='red'>{n}</font>", styleN))
        else:
            elements.append(Paragraph(n, styleN))
        elements.append(Spacer(1, 4))
        

# -----------------------------------------------------
# PREMIUM TWO-COLUMN BLOCK (TERMS LEFT, BANK RIGHT)
# -----------------------------------------------------

    partA = notes.get("partA", {})
    partB = notes.get("partB", [])

    # LEFT COLUMN — TERMS & CONDITIONS
    terms_title = Paragraph("<b>TERMS & CONDITIONS</b>", styleB)

    terms_data = [
       [Paragraph(f"<b>DELIVERY:</b> {partA.get('delivery','')}", styleN)],
       [Paragraph(f"<b>PERIOD:</b> {partA.get('period','')}", styleN)],
       [Paragraph(f"<b>TAX:</b> {partA.get('tax','')}", styleN)],
       [Paragraph(f"<b>P & F:</b> {partA.get('pf','')}", styleN)],
       [Paragraph(f"<b>PAYMENT:</b> {partA.get('payment','')}", styleN)],
       [Paragraph(f"<b>VALIDITY:</b> {partA.get('validity','')}", styleN)],
    ]

    terms_table = Table([[terms_title]] + terms_data, colWidths=[260])
    terms_table.setStyle(TableStyle([
    ('BOX', (0,0), (-1,-1), 1, colors.black),
    ('BACKGROUND', (0,0), (0,0), colors.whitesmoke),
    ('LEFTPADDING', (0,0), (-1,-1), 6),
    ('TOPPADDING', (0,0), (-1,-1), 5),
    ('BOTTOMPADDING', (0,0), (-1,-1), 5),
    ]))



# RIGHT COLUMN — BANK DETAILS
    bank_title = Paragraph("<b>BANK DETAILS</b>", styleB)

    bank_data = []

    BANK_DETAILS = [
    "<b>BANK NAME :</b> KOTAK MAHINDRA BANK.",
    "<b>ADD:</b> GROUND AND MEZZANINE FLOOR,",
    "BOTAWALA CHAMBER,2, MUMBAI- 01.",   # ← NOT BOLD
    "<b>BRANCH CODE :</b>957",
    "<b>IFSC CODE NO  :</b> KKBK0000957",
    "<b>OUR A/C NO :</b> 9223312803"
    ]

    for line in BANK_DETAILS:
        bank_data.append([Paragraph(line, styleN)])

    bank_table = Table(
        [[bank_title]] + bank_data,
        colWidths=[260]
    )

    bank_table.setStyle(TableStyle([
    ('BOX', (0,0), (-1,-1), 1, colors.black),
    ('BACKGROUND', (0,0), (0,0), colors.whitesmoke),
    ('LEFTPADDING', (0,0), (-1,-1), 6),
    ('TOPPADDING', (0,0), (-1,-1), 5),
    ('BOTTOMPADDING', (0,0), (-1,-1), 5),
    ]))

    # -------------------------------------
    # Create a temporary canvas for measurement
    temp_canvas = Canvas("temp.pdf", pagesize=letter)

# Measure heights correctly
    tw, th = terms_table.wrapOn(temp_canvas, 0, 0)
    bw, bh = bank_table.wrapOn(temp_canvas, 0, 0)

# Determine final equal height
    final_height = max(th, bh)

# Force equal-height rows
    terms_table._rowHeights = [final_height / len(terms_table._cellvalues)] * len(terms_table._cellvalues)
    bank_table._rowHeights  = [final_height / len(bank_table._cellvalues)] * len(bank_table._cellvalues)

# Combine side-by-side
    combined = Table(
       [[terms_table, bank_table]],
       colWidths=[270, 270]
    )
    combined.setStyle(TableStyle([
       ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))

    elements.append(combined)
    elements.append(Spacer(1, 20))
    # ----------------------------
    # FOOTER
    # ----------------------------
    
    footer_lines = "<br/>".join(notes["footer"])

    elements.append(Paragraph(footer_lines, styleN))
    elements.append(Spacer(1, 10))

    elements.append(Paragraph(
    "<para>Warm Regards,<br/><b>Sales Team</b><br/><b>CALIBER ENTERPRISE</b></para>",
    styleN
    ))

    elements.append(Spacer(1, 20))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_excel(quotation):
    output = BytesIO()
    
    # Create workbook
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        workbook  = writer.book
        sheet     = workbook.add_worksheet("Quotation")

        # Excel formats
        bold  = workbook.add_format({'bold': True})
        wrap  = workbook.add_format({'text_wrap': True})
        bold_wrap = workbook.add_format({'bold': True, 'text_wrap': True})

        row = 0
        # ---------------------------------------
        # COMPANY LOGO AT TOP
        # ---------------------------------------
        try:
           logo_path = "logo.jpeg"
           sheet.insert_image(0, 0, logo_path, {'x_scale': 0.7, 'y_scale': 0.7})
           row = 6  # move start down so content doesn't overlap logo
        except Exception as e:
          print("Excel logo error:", e)


        # -----------------------
        # HEADER
        # -----------------------
        sheet.write(row, 0, "QUOTATION", bold)
        row += 2

        for k, v in quotation["header"].items():
            sheet.write(row, 0, f"{k}:", bold)
            sheet.write(row, 1, str(v))
            row += 1

        row += 1

        # Intro
        sheet.write(row, 0, "Introduction:", bold)
        sheet.write(row+1, 0, quotation["intro"], wrap)
        row += 3

        # -----------------------
        # ITEMS TABLE
        # -----------------------
        headers = quotation["items_table"]["headers"]
        rows    = quotation["items_table"]["rows"]

        for col, h in enumerate(headers):
            sheet.write(row, col, h, bold_wrap)

        row += 1

        for r in rows:
            for col, cell in enumerate(r):
                sheet.write(row, col, str(cell), wrap)
            row += 1

        row += 2

        # -----------------------
        # NOTES
        # -----------------------
        notes = quotation["notes"]

        # Main Note
        sheet.write(row, 0, "Main Note:", bold)
        row += 1
        sheet.write(row, 0, notes["main_note"], wrap)
        row += 2

        # Sub Notes
        sheet.write(row, 0, "Sub Notes:", bold)
        row += 1

        for n in notes["sub_notes"]:
            sheet.write(row, 0, f"• {n}", wrap)
            row += 1

        row += 1

        # Part A
        sheet.write(row, 0, "Commercial Terms (Part A):", bold)
        row += 1

        for key, val in notes["partA"].items():
            sheet.write(row, 0, f"{key.capitalize()}:", bold)
            sheet.write(row, 1, str(val), wrap)
            row += 1

        row += 1

        # Part B
        sheet.write(row, 0, "Bank Details (Part B):", bold)
        row += 1

        for line in notes["partB"]:
            sheet.write(row, 0, f"• {line}", wrap)
            row += 1

        row += 2

        # Footer
       
        sheet.write(row, 0, "Footer:", bold)
        row += 1

        for line in notes["footer"]:
           sheet.write(row, 0, f"• {line}", wrap)
           row += 1


    output.seek(0)
    return output


def generate_word(quotation):
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO

    doc = Document()

    # ---------------------------------------
    # COMPANY LOGO AT TOP
    # ---------------------------------------
    try:
        logo_path = "logo.jpeg"
        doc.add_picture(logo_path, width=Inches(6.5))
    except Exception as e:
        print("Word logo error:", e)

    doc.add_heading("Quotation", level=1)

    # -----------------------------
    # HEADER
    # -----------------------------
    for k, v in quotation["header"].items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_paragraph("")  # spacing

    # -----------------------------
    # ITEMS TABLE
    # -----------------------------
    headers = quotation["items_table"]["headers"]
    rows = quotation["items_table"]["rows"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)

    # Data rows
    for row in rows:
        cells = table.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = str(c)

    doc.add_paragraph("")  # spacing

    # -----------------------------
    # NOTES SECTION
    # -----------------------------
    notes = quotation.get("notes", {})

    # MAIN NOTE
    if notes.get("main_note"):
        doc.add_heading("Main Note", level=2)
        doc.add_paragraph(notes["main_note"])

    # SUB NOTES
    if notes.get("sub_notes"):
        doc.add_heading("Sub Notes", level=3)
        sub = notes["sub_notes"]

        if isinstance(sub, list):
            for line in sub:
                if str(line).strip():
                    doc.add_paragraph(f"• {line}")

        elif isinstance(sub, str):
            for line in sub.split("\n"):
                if line.strip():
                    doc.add_paragraph(f"• {line}")

    # -----------------------------
    # PART A – COMMERCIAL TERMS
    # -----------------------------
    if notes.get("partA"):
        doc.add_heading("Commercial Terms (Part A)", level=3)
        partA = notes["partA"]

        if isinstance(partA, dict):
            for key, val in partA.items():
                doc.add_paragraph(f"{key.capitalize()}: {val}")

        elif isinstance(partA, list):
            for line in partA:
                doc.add_paragraph(str(line))

        else:
            for line in str(partA).split("\n"):
                doc.add_paragraph(line)

    # -----------------------------
    # PART B – BANK DETAILS
    # -----------------------------
    if notes.get("partB"):
        doc.add_heading("Bank Details (Part B)", level=3)
        partB = notes["partB"]

        if isinstance(partB, dict):
            for key, val in partB.items():
                doc.add_paragraph(f"{key.capitalize()}: {val}")

        elif isinstance(partB, list):
            for line in partB:
                doc.add_paragraph(str(line))

        else:
            for line in str(partB).split("\n"):
                doc.add_paragraph(line)

    # -----------------------------
    # FOOTER
    # -----------------------------
    if notes.get("footer"):
        doc.add_heading("Footer", level=3)
        footer = notes["footer"]

        if isinstance(footer, list):
            for line in footer:
                doc.add_paragraph(str(line))
        else:
            doc.add_paragraph(str(footer))

    # -----------------------------
    # RETURN WORD FILE
    # -----------------------------
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------
# Rate/Weight Functions
@st.cache_data
def load_catalogue(file):
    raw = json.load(file)
    normalized = []

    for entry in raw:
        # Case 1: Already structured
        if "screw_type" in entry:
            if "stud price" in entry.get("screw_type", "").lower():
                entry["screw_type"] = "Stud Screw"
                # Clean metric
                if "dimensions_in_metric" in entry:
                    cleaned = []
                    for dim in entry["dimensions_in_metric"]:
                        new_dia = {k.replace(" ", ""): v for k, v in dim.get("diameter", {}).items()}
                        cleaned.append({"length_mm": dim.get("length_mm"), "diameter": new_dia})
                    entry["dimensions_in_metric"] = cleaned
                # Clean inch
                if "dimensions_in_inches" in entry:
                    cleaned = []
                    for dim in entry["dimensions_in_inches"]:
                        new_dia = {k.strip(): v for k, v in dim.get("diameter", {}).items()}
                        cleaned.append({"length": dim.get("length"), "diameter": new_dia})
                    entry["dimensions_in_inches"] = cleaned
            normalized.append(entry)

        # Case 2: Flat (Rivet style)
        elif "title" in entry and ("dimensions_unit" in entry or "data" in entry):
            screw_type = entry["title"]
            dim_unit = entry.get("dimensions_unit", "").lower()
            unit = entry.get("approx_count_unit", "N/A")
            if "stud" in screw_type.lower():
                screw_type = "Stud Screw"
            new_entry = {"screw_type": screw_type, "standard": "N/A", "unit": unit,
                         "dimensions_in_metric": [], "dimensions_in_inches": []}
            if dim_unit.startswith("inch"):
                for d in entry.get("data", []):
                    new_entry["dimensions_in_inches"].append({
                        "length": d.get("length"),
                        "diameter": {k.replace("diameter_", "").strip(): v for k, v in d.items() if k.startswith("diameter_")}
                    })
            elif dim_unit.startswith("metric"):
                for d in entry.get("data", []):
                    diam_clean = {k.replace("diameter_", "").replace(" ", ""): v for k, v in d.items() if k.startswith("diameter_")}
                    new_entry["dimensions_in_metric"].append({"length_mm": d.get("length_mm"), "diameter": diam_clean})
            normalized.append(new_entry)

        # Case 3: Nut structures
        elif "approx_count_per_50_kgs" in entry or "hex_locknuts_bsw_bsf_approx_weight_per_100_pcs" in entry:
            screw_type = entry.get("title", "Nut")

            def looks_like_inches(s):
                s = str(s)
                return '"' in s or '/' in s or "ba" in s.lower()

            if "approx_count_per_50_kgs" in entry:
                for nut_type, data_list in entry["approx_count_per_50_kgs"].items():
                    ne = {"screw_type": f"{screw_type} - {nut_type}", "standard": "N/A",
                          "unit": "Approx. Count per 50 kgs.", "dimensions_in_metric": [], "dimensions_in_inches": []}
                    for d in data_list:
                        size, val = d.get("size"), d.get("specification") or d.get("height_d") or d.get("count")
                        target = "dimensions_in_inches" if looks_like_inches(size) else "dimensions_in_metric"
                        ne[target].append({"length_mm": None, "diameter": {size: val}})
                    normalized.append(ne)

            if "hex_locknuts_bsw_bsf_approx_weight_per_100_pcs" in entry:
                for nut_type, data_list in entry["hex_locknuts_bsw_bsf_approx_weight_per_100_pcs"].items():
                    ne = {"screw_type": f"{screw_type} - {nut_type}", "standard": "N/A",
                          "unit": "Approx. Weight per 100 pcs.", "dimensions_in_metric": [], "dimensions_in_inches": []}
                    for d in data_list:
                        size, val = d.get("size"), d.get("weight") or d.get("count")
                        target = "dimensions_in_inches" if looks_like_inches(size) else "dimensions_in_metric"
                        ne[target].append({"length_mm": None, "diameter": {size: val}})
                    normalized.append(ne)

        # Case 4: Washer structure
        elif "title" in entry and "unit" in entry:
            for ftype, data_list in entry.items():
                if ftype in ["title", "unit"]:
                    continue
                ne = {"screw_type": ftype, "standard": "N/A", "unit": entry.get("unit", "N/A"),
                      "dimensions_in_metric": [], "dimensions_in_inches": []}
                for d in data_list:
                    size, wt = d.get("size"), d.get("weight")
                    ne["dimensions_in_metric"].append({"length_mm": None, "diameter": {size: wt}})
                normalized.append(ne)
    return normalized


# -------------------------
# Utility Functions
def find_rate(entries, length_val, diameter, dim_type):
    dims_key = "dimensions_in_metric" if dim_type == "metric" else "dimensions_in_inches"
    for e in entries:
        for dim in e.get(dims_key, []):
            if dim.get("length_mm") == length_val or dim.get("length") == length_val or dim.get("length_mm") is None:
                val = dim.get("diameter", {}).get(diameter)
                if val:
                    return val
    return None


def sort_diameters(diam_list, dim_type):
    def metric_key(x):
        try:
            return int(''.join(ch for ch in x if ch.isdigit()) or 9999)
        except:
            return 9999
    def inch_key(x):
        try:
            val = x.replace('"', "")
            if "/" in val:
                n, d = val.split("/")
                return float(n) / float(d)
            return float(val)
        except:
            return 9999
    return sorted(list(dict.fromkeys(diam_list)), key=metric_key if dim_type == "metric" else inch_key)


def normalize_name(name):
    return name.split(" - ")[0].strip().lower()


def parse_combo_input(inp):
    words = re.findall(r"[A-Za-z]+|\d+", inp.lower())
    combo, qty = [], 1
    for w in words:
        if w.isdigit():
            qty = int(w)
        elif w in ["stud", "nut", "washer", "bolt", "screw"]:
            combo.append({"type": w, "count": qty})
            qty = 1
    return combo

def detect_dimension_unit(desc):
    d = str(desc).upper().replace("×", "X")
    if re.search(r"\bM\d+", d) or "MM" in d or " DIN" in d or " ISO" in d:
        return "metric"
    if '"' in d or "INCH" in d or "#" in d:
        return "inch"
    return "metric"

# Utility function for generating quotation number
def generate_quotation_no():
    today = datetime.date.today()
    month = today.strftime("%m")
    year = today.year

    # Determine financial year
    if today.month >= 4:  # April or later → FY starts this year
        fy_start = year % 100
        fy_end = (year + 1) % 100
    else:
        fy_start = (year - 1) % 100
        fy_end = year % 100

    fy_str = f"{fy_start:02d}-{fy_end:02d}"

    # Keep a counter file
    counter_file = "quotation_counter.json"
    if os.path.exists(counter_file):
        with open(counter_file, "r") as f:
            counters = json.load(f)
    else:
        counters = {}

    # Key = financial year
    if fy_str not in counters:
        counters[fy_str] = 0

    counters[fy_str] += 1
    serial = counters[fy_str]

    with open(counter_file, "w") as f:
        json.dump(counters, f)

    return f"CE/{month}/{serial:05d}/{fy_str}"

def generate_current_date():
    today = datetime.date.today()
    return "DATE : " + today.strftime("%d/%m/%Y")

# --------------------------------------------
# Streamlit UI with 6 Tabs
st.set_page_config(page_title="Quotation + Rate Generator", layout="wide")
st.title("⚙️ Quotation & Screw Rate/Weight Generator")

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📂 Upload File", "🧾 Header Details", "🧰 Items Table",
    "🗒️ Notes", "💰 Rate / Weight Calculator","📄 Generate Quotation"
])

# ----------------------------
# Tab 1-5: Quotation Logic
# ----------------------------
with tab1:
    uploaded_file = st.file_uploader("Upload Excel, PDF, or JSON file", type=["xlsx","xls","pdf","json"])
    if uploaded_file:
        ext = uploaded_file.name.split(".")[-1].lower()
        if ext in ["xlsx","xls"]:
            df = extract_excel(uploaded_file)
        elif ext=="pdf":
            df = extract_pdf(uploaded_file)
        elif ext=="json":
            df = extract_json(uploaded_file)
        else:
            st.error("Unsupported file type")
            st.stop()
        st.success("✅ File processed successfully. Move to next tab!")

# --------------------------------------------
# TAB 2: HEADER  (FINAL WORKING VERSION)
# --------------------------------------------
if "df" in locals():

    # --- Create header state only once ---
    if "header_data" not in st.session_state:
        st.session_state["header_data"] = {
            "quotation_no": "REF : " + generate_quotation_no(),
            "date": generate_current_date(),
            # Client basic
        "client_name": "",
        "client_address": "",
        "attention": "",

        # Additional client fields
        "mobile_no": "",
        "pan_no": "",
        "pincode": "",
        "state": "",
        "email_id": "",
        "gst_no": ""
        }

    with tab2:
        st.subheader("🧾 Edit Header")
        saved = load_saved_values() 
        col1, col2 = st.columns(2)

    # ---------------------------
    # LEFT COLUMN
    # ---------------------------
        with col1:
            st.session_state["header_data"]["quotation_no"] = st.text_input(
                "Quotation No.",
                value=st.session_state["header_data"]["quotation_no"],
                disabled=True
            )
            st.session_state["header_data"]["client_name"] = smart_input(
               "Client Name", "client_name", saved
            )


            st.session_state["header_data"]["attention"] = st.text_input(
                "Attention",
                value=st.session_state["header_data"]["attention"]
            )

    # ---------------------------
    # RIGHT COLUMN
    # ---------------------------
        with col2:
            st.session_state["header_data"]["date"] = st.text_input(
                "Date",
                 value=st.session_state["header_data"]["date"],
                 disabled=True
            )

            st.session_state["header_data"]["client_address"] = smart_input(
                "Client Address", "client_address",saved
            )
    # ---------------------------
    # ADDITIONAL CLIENT DETAILS (3×3 TABLE STYLE)
    # ---------------------------
        st.markdown("### Additional Client Details")

        colA, colB, colC = st.columns(3)

    # ROW 1
        with colA:
            st.session_state["header_data"]["mobile_no"] = smart_input(
               "Mobile No", "mobile_no",saved
            )
        with colB:
            st.session_state["header_data"]["pincode"] = smart_input(
                "Pincode", "pincode", saved
            )
        with colC:
            st.session_state["header_data"]["email_id"] = smart_input(
               "Email ID", "email_id", saved
            )
    # ROW 2
        colA2, colB2, colC2 = st.columns(3)

        with colA2:
            st.session_state["header_data"]["pan_no"] = smart_input(
               "PAN No", "pan_no", saved
            )
        with colB2:
            st.session_state["header_data"]["state"] = smart_input(
                "State", "state", saved
            )
        with colC2:
            st.session_state["header_data"]["gst_no"] = smart_input(
                "GST No", "gst_no", saved
            )
    # -------------------------
    # TAB 3: REPLACED — Item Mapping (full app3.py logic integrated)
    # -------------------------
    with tab3:
        # --- Mapping UI & logic from app3.py integrated here ---
        st.subheader("🧰 Item Mapping & Material Standards (Auto-mapping from app3.py)")
        st.markdown("<small>Automatic mapping of Type, Standard, Grade, Finish, and Unit. Edit any dropdowns if required.</small>", unsafe_allow_html=True)

        # Load synonyms file (if present)
        SYNONYMS_FILENAME = "synonyms.json"
        try:
            with open(SYNONYMS_FILENAME, "r", encoding="utf-8") as f:
                synonyms_data = json.load(f)
        except Exception:
            synonyms_data = {}

        # Helper functions copied/adapted from app3.py
        def normalize(text):
            if text is None:
                return ""
            return re.sub(r"[^A-Z0-9]", "", str(text).upper())

        def normalize_preserve_space(text):
            if text is None:
                return ""
            return re.sub(r"[^A-Z0-9\s]", " ", str(text).upper()).strip()

        def normalized_token_in_text(token, text):
            if not token or not text:
                return False
            tnorm = normalize(str(token))
            if not tnorm:
                return False
            txt_norm = normalize(str(text))
            if tnorm in txt_norm:
                return True
            token_ws = normalize_preserve_space(str(token))
            text_ws = normalize_preserve_space(str(text))
            try:
                pattern = r"\b" + re.escape(token_ws) + r"\b"
                if re.search(pattern, text_ws):
                    return True
            except re.error:
                pass
            return False

        def parse_grades_field(g):
            if g is None:
                return []
            out = []
            if isinstance(g, list):
                for el in g:
                    if el is None:
                        continue
                    if isinstance(el, (int, float)):
                        out.append(str(el))
                    elif isinstance(el, str):
                        parts = re.split(r'[;,]', el)
                        out.extend([p.strip() for p in parts if p.strip()])
                    else:
                        out.append(str(el).strip())
                return out
            if isinstance(g, (int, float)):
                return [str(g)]
            if isinstance(g, str):
                parts = re.split(r'[;,]', g)
                return [p.strip() for p in parts if p.strip()]
            return []

        def split_comma_values_list(lst):
            out = []
            for v in lst:
                if isinstance(v, str):
                    if v.strip().startswith("[") and v.strip().endswith("]"):
                        try:
                            inner = json.loads(v.replace("'", '"'))
                            if isinstance(inner, list):
                                out.extend([str(x).strip() for x in inner if x])
                                continue
                        except Exception:
                            pass
                    parts = [p.strip() for p in re.split(r'[;,]', v) if p.strip()]
                    out.extend(parts)
                elif isinstance(v, list):
                    out.extend(split_comma_values_list(v))
                else:
                    out.append(v)
            return out

        def clean_din_json(data):
            if isinstance(data, dict):
                new_data = {}
                for k, v in data.items():
                    if isinstance(v, list):
                        new_data[k] = split_comma_values_list(v)
                    elif isinstance(v, dict):
                        new_data[k] = clean_din_json(v)
                    elif isinstance(v, str):
                        if "," in v or ";":
                            parts = [p.strip() for p in re.split(r'[;,]', v) if p.strip()]
                            new_data[k] = parts if len(parts) > 1 else parts[0]
                        else:
                            new_data[k] = v.strip()
                    else:
                        new_data[k] = v
                return new_data
            elif isinstance(data, list):
                return [clean_din_json(x) if isinstance(x, (dict, list)) else x for x in data]
            return data

        # Build din_raw from your existing din_data (loaded earlier) — but clean it first
        try:
            din_raw = clean_din_json(din_data)
        except Exception:
            din_raw = din_data

        # Build DIN index & finishes (from app3.py)
        def build_din_index(din_data):
            categories = {}
            global_finishes = []

            def add_item(cat_key, item):
                categories.setdefault(cat_key, []).append(item)

            if isinstance(din_data, dict):
                for topk, val in din_data.items():
                    topk_lower = str(topk).strip().lower()
                    if "finish" in topk_lower and isinstance(val, list):
                        for v in val:
                            if isinstance(v, str):
                                global_finishes.append(v.strip().upper())
                            elif isinstance(v, list):
                                for x in v:
                                    global_finishes.append(str(x).strip().upper())
                        continue
                    if isinstance(val, list):
                        category_key = topk_lower.rstrip('s')
                        for entry in val:
                            if not isinstance(entry, dict):
                                continue
                            type_key = None
                            for k in entry.keys():
                                if k.strip().lower().endswith(" type"):
                                    type_key = k
                                    break
                            type_name = entry.get(type_key) if type_key else entry.get("Type") or entry.get("Bolt Type") or entry.get("Nut Type") or entry.get("name") or ""
                            standard = entry.get("Standard") or entry.get("standard") or ""
                            inches = entry.get("Inches") or entry.get("inches") or ""
                            metrics = entry.get("Metrics") or entry.get("metrics") or ""
                            grades = parse_grades_field(entry.get("Grades") or entry.get("grades"))
                            entry_finishes = []
                            for k, v in entry.items():
                                if "finish" in k.lower():
                                    if isinstance(v, list):
                                        for el in v:
                                            if el:
                                                entry_finishes.append(str(el).strip())
                                    elif isinstance(v, str):
                                        entry_finishes.extend([x.strip() for x in re.split(r'[;,]', v) if x.strip()])
                            add_item(category_key, {
                                "type_name": str(type_name).strip(),
                                "standard": str(standard).strip(),
                                "inches": inches,
                                "metrics": metrics,
                                "grades": grades,
                                "finishes": [f for f in entry_finishes],
                                "raw": entry
                            })
                            for f in entry_finishes:
                                if f:
                                    global_finishes.append(str(f).upper())
            elif isinstance(din_data, list):
                for entry in din_data:
                    if not isinstance(entry, dict):
                        continue
                    type_key = None
                    for k in entry.keys():
                        if k.strip().lower().endswith(" type"):
                            type_key = k
                            break
                    if type_key:
                        category_key = type_key.strip().split()[0].lower()
                    else:
                        category_key = "unknown"
                        for guess in ["washer", "bolt", "nut", "stud", "screw"]:
                            if any(guess in k.lower() for k in entry.keys()):
                                category_key = guess
                                break
                    type_name = entry.get(type_key) if type_key else entry.get("Type") or entry.get("Bolt Type") or entry.get("Nut Type") or entry.get("name") or ""
                    standard = entry.get("Standard") or entry.get("standard") or ""
                    inches = entry.get("Inches") or entry.get("inches") or ""

                    raw_metrics = entry.get("Metrics") or entry.get("metrics") or []
                    if isinstance(raw_metrics, list):
                       metrics = []
                       for m in raw_metrics:
                            if isinstance(m, str):
                               parts = [p.strip() for p in re.split(r'[;,]', m) if p.strip()]
                               metrics.extend(parts)
                            else:
                               metrics.append(str(m).strip())
                    else:
                        metrics = [p.strip() for p in re.split(r'[;,]', str(raw_metrics)) if p.strip()]

                    grades = parse_grades_field(entry.get("Grades") or entry.get("grades"))

                    entry_finishes = []
                    for k, v in entry.items():
                        if "finish" in k.lower():
                            if isinstance(v, list):
                                for el in v:
                                    if el:
                                        entry_finishes.append(str(el).strip())
                            elif isinstance(v, str):
                                entry_finishes.extend([x.strip() for x in re.split(r'[;,]', v) if x.strip()])
                    add_item(category_key, {
                        "type_name": str(type_name).strip(),
                        "standard": str(standard).strip(),
                        "inches": inches,
                        "metrics": metrics,
                        "grades": grades,
                        "finishes": [f for f in entry_finishes],
                        "raw": entry
                    })
                    for f in entry_finishes:
                        if f:
                            global_finishes.append(str(f).upper())

            seen = set()
            dedup_fin = []
            for f in global_finishes:
                fu = str(f).upper().strip()
                if fu and fu not in seen:
                    seen.add(fu)
                    dedup_fin.append(fu)
            return categories, dedup_fin

        # Build index
        din_index, din_global_finishes = build_din_index(din_raw)

        # Build type maps
        type_to_entries = {}
        for cat, items in din_index.items():
            for it in items:
                tname = (it.get("type_name") or "").strip()
                if not tname:
                    continue
                type_to_entries.setdefault(tname, []).append(it)
        category_to_types = {cat: sorted(list({it['type_name'] for it in items if it.get('type_name')})) for cat, items in din_index.items()}

        # Candidate extraction and grade detection functions (adapted)
        def extract_candidate_terms(desc):
            patterns = [
                r"SA\d{3,4}\s*GR\s*[A-Z0-9]+",
                r"A\d{3,4}\s*[A-Z0-9]+",
                r"\b[A-Z]{1,3}\d{2,4}[A-Z0-9\-]*\b",
                r"HASTELLOY\s*[A-Z0-9]*",
                r"INCONEL\s*[A-Z0-9]*",
                r"ALLOY\s*[0-9A-Z\-]*",
                r"MONEL\s*[0-9A-Z]*",
                r"TI[A-Z0-9.\-]+",
                r"F\d{2,3}",
                r"\bCL\s*\d+(\.\d+)?\b",
                r"\b\d+\.\d+\b",
                r"ASTM\s*[A-Z]?\d{1,4}[A-Z0-9\-]*"
            ]
            found = []
            desc_upper = desc.upper()
            for pat in patterns:
                for match in re.findall(pat, desc_upper):
                    if match and match not in found:
                        found.append(match.strip())
            return found

        def extract_all_grades(obj, collected=None, title=None):
            if collected is None:
                collected = []
            if isinstance(obj, dict):
                local_title = title or obj.get("title", "")
                for k, v in obj.items():
                    if isinstance(v, (dict, list)):
                        extract_all_grades(v, collected, local_title)
                    elif isinstance(v, (str, int, float)):
                        v_clean = str(v).strip()
                        if v_clean and len(v_clean) > 0:
                            collected.append((v_clean, local_title, k))
            elif isinstance(obj, list):
                for item in obj:
                    extract_all_grades(item, collected, title)
            return collected

        all_grades = extract_all_grades(material_data)

        def get_grades_from_desc(desc, threshold=85):
            desc_upper = desc.upper()
            candidates = extract_candidate_terms(desc)
            matches = []
            exact_match = None
            for cand in candidates:
                norm_cand = normalize(cand)
                for grade_val, title, key in all_grades:
                    norm_grade = normalize(str(grade_val))
                    if not norm_grade:
                        continue
                    if norm_grade in normalize(desc_upper):
                        exact_match = grade_val
                    if norm_grade in norm_cand or norm_cand in norm_grade:
                        matches.append((grade_val, title, key))
                    elif fuzz.partial_ratio(norm_grade, norm_cand) >= threshold:
                        matches.append((grade_val, title, key))
            if not matches:
                for grade_val, title, key in all_grades:
                    if fuzz.partial_ratio(normalize(str(grade_val)), desc_upper) >= threshold:
                        matches.append((grade_val, title, key))
            if exact_match:
                matches = sorted(matches, key=lambda x: 0 if x[0] == exact_match else 1)
            seen = set()
            out = []
            for g, t, k in matches:
                if g not in seen:
                    seen.add(g)
                    out.append(g)
            return out

        def get_finish_from_desc(desc, threshold=85):
            desc_upper = " " + re.sub(r"[^A-Z0-9\s]", " ", (desc or "").upper()) + " "
            matched_finishes = []
            for finish in din_global_finishes:
                fin_norm = finish.upper().strip()
                if not fin_norm:
                    continue
                try:
                    if re.search(r"\b" + re.escape(fin_norm) + r"\b", desc_upper):
                        matched_finishes.append(fin_norm)
                    elif fuzz.partial_ratio(fin_norm, desc_upper) >= threshold:
                        matched_finishes.append(fin_norm)
                except re.error:
                    if fin_norm in desc_upper:
                        matched_finishes.append(fin_norm)
            for tok in ["ZINC", "HDG", "BLACK", "PLATED", "GALVANIZED", "PASSIVATED"]:
                if tok in desc_upper and tok not in matched_finishes:
                    matched_finishes.append(tok)
            return matched_finishes

        # Category detection
        CATEGORY_KEYWORDS = {
            "stud": ["STUD", "STUDS"],
            "nut": ["NUT", "NUTS"],
            "bolt": ["BOLT", "BOLTS", "CAPSCREW", "HEXHD", "HEX HEAD", "SOCKET", "GRUB"],
            "washer": ["WASHER", "WASHERS", "FLAT WASHER", "SPRING WASHER", "TOOTH"],
            "screw": ["SCREW", "SCREWS"]
        }
        def detect_category(desc):
            desc_norm = normalize_preserve_space(desc)
            detected = []
            for cat, keywords in CATEGORY_KEYWORDS.items():
                for kw in keywords:
                    if kw in desc_norm:
                        detected.append(cat)
                        break
            for cat in din_index.keys():
                if cat and cat in desc_norm and cat not in detected:
                    detected.append(cat)
            seen = set()
            out = []
            for d in detected:
                if d not in seen:
                    out.append(d)
                    seen.add(d)
            return out

        # Synonym match function (from app3)
        def find_all_synonym_matches(desc, synonyms_dict):
            if not desc:
                return []
            desc_up = desc.upper()
            out = []

            if "CAPSCREW" in desc_up and ("HEX HD" in desc_up or "HEXHD" in desc_up or "HEX HEAD" in desc_up):
                for main, syns in synonyms_dict.items():
                    syn_list = syns if isinstance(syns, list) else [syns]
                    for s in syn_list:
                        if s and s.upper() in desc_up:
                            if main.upper() in ["ALLEN CAP SCREW", "HEX BOLT"]:
                                out.append((s, main))
                return out

            for main, syns in synonyms_dict.items():
                syn_list = syns if isinstance(syns, list) else [syns]
                for s in syn_list:
                    if s and s.upper() in desc_up:
                        tok = s.upper()
                        if ("CAPSCREW" in tok or "ALLEN CAP" in tok or "SOCKET CAP" in tok or "ALLEN" in tok) and main.upper() != "ALLEN CAP SCREW":
                            continue
                        out.append((s, main))
            return out

        # UI styling
        st.markdown(
            """
            <style>
            div[role="listbox"], div[role="combobox"], div.stSelectbox > div[role="button"] {
                min-width: 320px !important;
                max-width: 100% !important;
                white-space: normal !important;
                overflow: visible !important;
            }
            .stSelectbox [role="listbox"] { min-width: 420px !important; white-space: normal !important; }
            .row-card { border: 1px solid rgba(0,0,0,0.06); padding: 8px; border-radius: 8px; margin-bottom: 6px; background: #fff; }
            .row-header { background: rgba(0,0,0,0.03); padding: 6px; border-radius: 6px; margin-bottom: 8px; }
            .small-selected { font-size:12px; color:#666; margin-bottom:4px; min-height:14px; }
            </style>
            """,
            unsafe_allow_html=True,
        )

        # Find description column heuristically (use same heuristics as app3)
        desc_col = None
        for col in df.columns:
            col_lower = str(col).lower()
            if any(word in col_lower for word in ["desc", "description", "material", "details", "part", "combined", "extract"]):
                desc_col = col
                break
        if not desc_col:
            for col in df.columns:
                sample = " ".join(df[col].astype(str).head(5).tolist()).lower()
                if any(word in sample for word in ["bolt", "nut", "screw", "washer", "steel", "stud", "thread"]):
                    desc_col = col
                    break
        if not desc_col:
            desc_col = df.columns[0]

        # ensure required columns exist in df
        for c in ["Grade", "Finish", "DIN Category", "Type", "Standard", "Inches", "Metrics", "DIN Details", "Detected Unit", "Standard Family"]:
            if c not in df.columns:
                df[c] = ""

        # Precompute suggestions
        grade_results = []
        finish_suggestions = []
        category_suggestions = []
        unit_suggestions = []
        for _, row in df.iterrows():
            desc = str(row[desc_col])
            grade_results.append(get_grades_from_desc(desc))
            finish_suggestions.append(get_finish_from_desc(desc))
            category_suggestions.append(detect_category(desc))
            unit_suggestions.append(detect_dimension_unit(desc) if 'detect_dimension_unit' in globals() else ("metric" if "M" in str(desc).upper() or "MM" in str(desc).upper() else "inch"))

        st.markdown("<div class='row-header'><b>Mapped rows — editable</b></div>", unsafe_allow_html=True)

        # layout widths
        block_widths = [0.3, 3.2, 1.6, 1.0, 0.9, 2.0, 1.4, 1.6]

        st_columns = st.columns([0.3, 3.2])
        st_columns[0].markdown("**#**")
        st_columns[1].markdown("**Description (mapped types below)**")

        # Render mapping UI row-by-row (adapted from app3)
        for i, row in df.iterrows():
            orig_desc = str(row[desc_col])
            matched_syn_pairs = find_all_synonym_matches(orig_desc, synonyms_data)
            matched_syn = matched_syn_pairs[0][0] if matched_syn_pairs else None
            matched_mains = [m for s, m in matched_syn_pairs]

            preferred_types_from_syn = {'nut': [], 'bolt': [], 'screw': [], 'washer': [], 'stud': []}
            if matched_mains:
                for main in matched_mains:
                    m_up = str(main).upper()
                    if "NUT" in m_up:
                        preferred_types_from_syn['nut'].append(str(main).strip())
                    elif "BOLT" in m_up:
                        preferred_types_from_syn['bolt'].append(str(main).strip())
                    elif "SCREW" in m_up:
                        preferred_types_from_syn['screw'].append(str(main).strip())
                    elif "WASHER" in m_up:
                        preferred_types_from_syn['washer'].append(str(main).strip())
                    elif "STUD" in m_up:
                        preferred_types_from_syn['stud'].append(str(main).strip())
                    else:
                        for cat, opts in category_to_types.items():
                            for opt in opts:
                                if normalize(opt) == normalize(main):
                                    if cat in preferred_types_from_syn:
                                        preferred_types_from_syn[cat].append(str(main).strip())

            if matched_syn_pairs:
                desc = orig_desc
                desc_for_match = orig_desc
                for s, m in matched_syn_pairs:
                    try:
                        desc_for_match = re.sub(re.escape(s), str(m), str(desc_for_match), flags=re.IGNORECASE)
                    except Exception:
                        desc_for_match = desc_for_match.replace(s, m)
                for main_category in matched_mains:
                    m_up = str(main_category).upper()
                    if "NUT" in m_up and "nut" not in category_suggestions[i]:
                        category_suggestions[i].append("nut")
                    if "BOLT" in m_up and "bolt" not in category_suggestions[i]:
                        category_suggestions[i].append("bolt")
                    if "SCREW" in m_up and "screw" not in category_suggestions[i]:
                        category_suggestions[i].append("screw")
            else:
                desc = orig_desc
                desc_for_match = orig_desc

            inferred_cats = []
            for m in matched_mains:
                m_up = str(m).upper()
                if "NUT" in m_up and "nut" not in inferred_cats:
                    inferred_cats.append("nut")
                if "BOLT" in m_up and "bolt" not in inferred_cats:
                    inferred_cats.append("bolt")
                if "SCREW" in m_up and "screw" not in inferred_cats:
                    inferred_cats.append("screw")
                if "WASHER" in m_up and "washer" not in inferred_cats:
                    inferred_cats.append("washer")
                if "STUD" in m_up and "stud" not in inferred_cats:
                    inferred_cats.append("stud")

            detected_cats = list(dict.fromkeys((category_suggestions[i] or []) + inferred_cats))
            if not detected_cats:
                detected_cats = [""]

            for idx in range(len(detected_cats)):
                tcol = f"Type_{idx+1}"
                scol = f"Standard_{idx+1}"
                gcol = f"Grade_{idx+1}"
                fcol = f"Finish_{idx+1}"
                ucol = f"Unit_{idx+1}"
                famcol = f"Standard_Family_{idx+1}"
                if tcol not in df.columns:
                    df[tcol] = ""
                if scol not in df.columns:
                    df[scol] = ""
                if gcol not in df.columns:
                    df[gcol] = ""
                if fcol not in df.columns:
                    df[fcol] = ""
                if ucol not in df.columns:
                    df[ucol] = ""
                if famcol not in df.columns:
                    df[famcol] = ""

            first_block = True
            for t_idx, cat in enumerate(detected_cats):
                cnum, cdesc, ctype_col, cfamily_col, cunit_col, cdim_col, cgrade_col, cfinish_col = st.columns(block_widths)

                if first_block:
                    cnum.markdown(f"<div style='padding:6px;border:1px solid rgba(0,0,0,0.06);border-radius:6px;text-align:center'>{i+1}</div>", unsafe_allow_html=True)
                    cdesc.markdown(f"<div style='padding:6px;border:1px solid rgba(0,0,0,0.06);border-radius:6px'>{(desc if len(desc) < 300 else desc[:297] + '...')}</div>", unsafe_allow_html=True)
                else:
                    cnum.markdown("<div style='height:40px'></div>", unsafe_allow_html=True)
                    cdesc.markdown("<div style='height:40px'></div>", unsafe_allow_html=True)

                if first_block and matched_syn_pairs:
                    note_parts = []
                    seen_notes = set()
                    for s, m in matched_syn_pairs:
                        label = f"({s} is synonym of {m})"
                        if label not in seen_notes:
                            note_parts.append(label)
                            seen_notes.add(label)
                    if note_parts:
                        note_html = "<div style='font-size:12px;color:gray;margin-top:2px;'>" + " ".join(note_parts) + "</div>"
                        cdesc.markdown(note_html, unsafe_allow_html=True)

                key_t = f"row_{i}_type_{t_idx+1}"
                key_fam = f"row_{i}_family_{t_idx+1}"
                key_unit = f"row_{i}_unit_{t_idx+1}"
                key_dim = f"row_{i}_dim_{t_idx+1}"
                key_grade = f"row_{i}_grade_{t_idx+1}"
                key_finish = f"row_{i}_finish_{t_idx+1}"

                desc_up = desc_for_match.upper()

                t_opts = []
                if cat:
                    t_opts = category_to_types.get(cat, []) or []
                if not t_opts:
                    all_types = []
                    for k, v in category_to_types.items():
                        all_types.extend(v)
                    seent = set()
                    t_opts = [x for x in all_types if x and (x not in seent and not seent.add(x))]

                pref_for_cat = None
                if cat and preferred_types_from_syn.get(cat):
                    for candidate in preferred_types_from_syn.get(cat):
                        match_opt = next((o for o in t_opts if normalize(o) == normalize(candidate)), None)
                        if match_opt:
                            pref_for_cat = match_opt
                            break
                    for candidate in preferred_types_from_syn.get(cat):
                        if candidate and all(normalize(candidate) != normalize(o) for o in t_opts):
                            t_opts.insert(0, candidate)

                existing_type_val = ""
                col_type_name = f"Type_{t_idx+1}"
                if df.at[i, col_type_name]:
                    existing_type_val = df.at[i, col_type_name]

                pre_type = existing_type_val or pref_for_cat or ""
                if not pre_type and cat:
                    if cat == "bolt":
                        pre_type = next((opt for opt in t_opts if opt.upper() == "HEX BOLT"), "")
                    elif cat == "nut":
                        pre_type = next((opt for opt in t_opts if opt.upper() == "HEX NUT"), "")
                if not pre_type and t_opts:
                    pre_type = t_opts[0]

                display_type = st.session_state.get(key_t, pre_type)
                ctype_col.markdown(f"<div class='small-selected'>{display_type}</div>", unsafe_allow_html=True)
                try:
                    t_options_final = [""] + t_opts
                    t_index = t_options_final.index(display_type) if display_type in t_options_final else (t_options_final.index(pre_type) if pre_type in t_options_final else 0)
                except Exception:
                    t_index = 0
                sel_type = ctype_col.selectbox("", options=t_options_final, index=t_index, key=key_t, help=f"Type (category: {cat or 'any'})")

                chosen_type = st.session_state.get(key_t, sel_type) or ""
                df.at[i, col_type_name] = chosen_type

                # Family options
                fams_all = set()
                for items in din_index.values():
                    for it in items:
                        sv = str(it.get("standard") or "").strip()
                        fam = sv.split()[0] if sv else ""
                        if fam:
                            fams_all.add(fam)
                sel_types_for_family = [chosen_type] if chosen_type else []
                fams_for_selected_types = set()
                for tsel in sel_types_for_family:
                    for items in din_index.values():
                        for it in items:
                            if it.get("type_name") == tsel:
                                sv = str(it.get("standard") or "").strip()
                                fam = sv.split()[0] if sv else ""
                                if fam:
                                    fams_for_selected_types.add(fam)
                if len(sel_types_for_family) == 1 and fams_for_selected_types:
                    family_options = [""] + sorted(list(fams_for_selected_types))
                else:
                    family_options = [""] + sorted(list(fams_all))

                prev_family = df.at[i, "Standard"] and (df.at[i, "Standard"].split()[0] if df.at[i, "Standard"] else "")
                fam_idx = 0
                if prev_family and prev_family in family_options:
                    fam_idx = family_options.index(prev_family)
                else:
                    for f in family_options:
                        if f and f in desc_for_match.upper():
                            fam_idx = family_options.index(f)
                            break
                display_family = st.session_state.get(key_fam, family_options[fam_idx] if fam_idx < len(family_options) else "")
                cfamily_col.markdown(f"<div class='small-selected'>{display_family}</div>", unsafe_allow_html=True)
                family_sel = cfamily_col.selectbox("", options=family_options, index=fam_idx, key=key_fam, help="Standard family (DIN/ISO/ASME etc.)")
                df.at[i, f"Standard_Family_{t_idx+1}"] = st.session_state.get(key_fam, family_sel) or ""

                # Unit
                unit_options = ["metric", "inch"]
                auto_unit = unit_suggestions[i] if i < len(unit_suggestions) else ""
                default_unit = st.session_state.get(key_unit, df.at[i, f"Unit_{t_idx+1}"] if (ucol := f"Unit_{t_idx+1}") in df.columns else "") or (auto_unit if auto_unit in ("metric","inch") else "")
                if not default_unit:
                    default_unit = "metric" if "M" in desc_for_match.upper() or "MM" in desc_for_match.upper() else ("inch" if "#" in desc_for_match or "INCH" in desc_for_match.upper() else "metric")
                display_unit = st.session_state.get(key_unit, default_unit)
                cunit_col.markdown(f"<div class='small-selected'>{display_unit}</div>", unsafe_allow_html=True)
                try:
                    unit_index = unit_options.index(default_unit) if default_unit in unit_options else 0
                except Exception:
                    unit_index = 0
                unit_sel = cunit_col.selectbox("", options=unit_options, index=unit_index, key=key_unit, help="Unit (metric/inch)")
                df.at[i, f"Unit_{t_idx+1}"] = st.session_state.get(key_unit, unit_sel) or unit_sel

                # Dimensions (filtered)
                dim_opts = []
                current_sel = chosen_type or ""
                current_family = st.session_state.get(key_fam, df.at[i, f"Standard_Family_{t_idx+1}"] or "") or ""
                current_unit = st.session_state.get(key_unit, df.at[i, f"Unit_{t_idx+1}"] or "") or ""

                for catk, items in din_index.items():
                    for it in items:
                        try:
                            if not it.get("type_name"):
                                continue
                            if current_sel and it.get("type_name").strip() != current_sel.strip():
                                continue
                            std_raw = str(it.get("standard") or "").strip()
                            metrics_raw = it.get("metrics") or ""
                            if isinstance(metrics_raw, list):
                                metrics_raw = ", ".join([str(m).strip() for m in metrics_raw if m])
                            inches_raw = str(it.get("inches") or "").strip()
                            if current_family:
                                fam_it = std_raw.split()[0] if std_raw else ""
                                if fam_it != current_family:
                                    if std_raw and re.match(r"^\d", std_raw) and not metrics_raw:
                                        cand = f"{current_family} {std_raw}"
                                        dim_opts.append(cand)
                                    else:
                                        continue
                            if current_unit == "metric":
                                if metrics_raw:
                                    parts = [p.strip() for p in re.split(r'[;,]', str(metrics_raw)) if p.strip()]
                                    for p in parts:
                                        if p not in dim_opts:
                                            dim_opts.append(p)
                                else:
                                    if std_raw:
                                        if re.search(r"\bDIN\b|\bISO\b|\d", std_raw):
                                            if std_raw not in dim_opts:
                                                dim_opts.append(std_raw)
                            elif current_unit == "inch":
                                if inches_raw:
                                    if inches_raw not in dim_opts:
                                        dim_opts.append(inches_raw)
                                else:
                                    if std_raw:
                                        if re.search(r"INCH", std_raw.upper()) or re.search(r"['\"]", std_raw):
                                            if std_raw not in dim_opts:
                                                dim_opts.append(std_raw)
                            else:
                                if metrics_raw:
                                    parts = [p.strip() for p in re.split(r'[;,]', str(metrics_raw)) if p.strip()]
                                    for p in parts:
                                        if p not in dim_opts:
                                            dim_opts.append(p)
                                elif std_raw:
                                    if std_raw not in dim_opts:
                                        dim_opts.append(std_raw)
                        except Exception:
                            continue

                if not dim_opts and current_family:
                    for catk, items in din_index.items():
                        for it in items:
                            std_raw = str(it.get("standard") or "").strip()
                            metrics_raw = it.get("metrics") or ""
                            if isinstance(metrics_raw, list):
                                metrics_raw = ", ".join([str(m).strip() for m in metrics_raw if m])
                            inches_raw = str(it.get("inches") or "").strip()
                            fam_it = std_raw.split()[0] if std_raw else ""
                            if fam_it == current_family:
                                if current_unit == "metric":
                                    if metrics_raw and metrics_raw not in dim_opts:
                                        parts = [p.strip() for p in re.split(r'[;,]', str(metrics_raw)) if p.strip()]
                                        for p in parts:
                                            if p not in dim_opts:
                                                dim_opts.append(p)
                                elif current_unit == "inch":
                                    if inches_raw and inches_raw not in dim_opts:
                                        dim_opts.append(inches_raw)
                                else:
                                    cand = metrics_raw or std_raw or inches_raw
                                    if cand and cand not in dim_opts:
                                        dim_opts.append(cand)

                if not dim_opts:
                    seen_dim = set()
                    for catk, items in din_index.items():
                        for it in items:
                            metrics_raw = it.get("metrics") or ""
                            if isinstance(metrics_raw, list):
                                metrics_raw = ", ".join([str(m).strip() for m in metrics_raw if m])
                            std_raw = str(it.get("standard") or "").strip()
                            inches_raw = str(it.get("inches") or "").strip()
                            if current_unit == "metric" and metrics_raw:
                                parts = [p.strip() for p in re.split(r'[;,]', str(metrics_raw)) if p.strip()]
                                for p in parts:
                                    if p and p not in seen_dim:
                                        seen_dim.add(p)
                                        dim_opts.append(p)
                            elif current_unit == "inch" and inches_raw:
                                if inches_raw and inches_raw not in seen_dim:
                                    seen_dim.add(inches_raw)
                                    dim_opts.append(inches_raw)
                            else:
                                cand_list = [metrics_raw, std_raw, inches_raw]
                                for cand in cand_list:
                                    if cand and cand not in seen_dim:
                                        seen_dim.add(cand)
                                        dim_opts.append(cand)

                dim_list = []
                seen_dim2 = set()
                for d in dim_opts:
                    if not d:
                        continue
                    dclean = str(d).strip()
                    if dclean not in seen_dim2:
                        seen_dim2.add(dclean)
                        dim_list.append(dclean)

                dim_key = key_dim
                default_dim = df.at[i, f"Standard_{t_idx+1}"] or (dim_list[0] if dim_list else "")

                if matched_syn_pairs:
                    for cand in dim_list:
                        try:
                            if normalize(str(cand)) == normalize(str(matched_mains[0])):
                                default_dim = cand
                                break
                        except Exception:
                            continue

                display_dim = st.session_state.get(dim_key, default_dim)
                cdim_col.markdown(f"<div class='small-selected'>{display_dim}</div>", unsafe_allow_html=True)
                def options_and_normalized_index(options_list, desired_value):
                    opts = [""] + list(options_list)
                    if desired_value is None:
                        return opts, 0
                    try:
                        if desired_value in opts:
                            return opts, opts.index(desired_value)
                    except Exception:
                        pass
                    dnorm = normalize(str(desired_value))
                    for idx, o in enumerate(opts):
                        if not o:
                            continue
                        if normalize(str(o)) == dnorm:
                            return opts, idx
                    for idx, o in enumerate(opts):
                        if not o:
                            continue
                        if dnorm and dnorm in normalize(str(o)):
                            return opts, idx
                    return opts, 0
                dim_options_final, dim_idx = options_and_normalized_index(dim_list, default_dim)
                dim_sel = cdim_col.selectbox("", options=dim_options_final, index=dim_idx, key=dim_key, help="Dim Standard (filtered by Type + Family + Unit)")
                df.at[i, f"Standard_{t_idx+1}"] = st.session_state.get(dim_key, dim_sel) or ""

                # Grade
                grade_set = []
                for catk, items in din_index.items():
                    for it in items:
        # Match selected type with normalization
                        if normalize(str(it.get("type_name"))) != normalize(str(chosen_type)):
                           continue

        # Directly load all grades associated with this type
                        for g in it.get("grades", []):
                            if g:
                                gs = str(g).strip()
                                if gs not in grade_set:
                                   grade_set.append(gs)
                for g in grade_results[i]:
                    if g and g not in grade_set:
                        grade_set.insert(0, g)
                if not grade_set and all_grades:
                    for gval, _, _ in all_grades:
                        if gval not in grade_set:
                            grade_set.append(gval)

                default_grade = df.at[i, f"Grade_{t_idx+1}"] or ""
                exact_found = ""
                try:
                    for g in grade_set:
                        if normalized_token_in_text(g, desc_for_match):
                            exact_found = g
                            break
                    if not exact_found:
                        for gval, _, _ in all_grades:
                            if normalized_token_in_text(gval, desc_for_match):
                                exact_found = gval
                                break
                except Exception:
                    exact_found = ""
                manual_grade = st.session_state.get(key_grade, "")
                if exact_found and not manual_grade:
                   df.at[i, f"Grade_{t_idx+1}"] = exact_found
                   if key_grade not in st.session_state:
                      st.session_state[key_grade] = exact_found
                      default_grade = exact_found
                else:
                  default_grade = manual_grade or df.at[i, f"Grade_{t_idx+1}"] or (grade_set[0] if grade_set else "")

                display_grade = st.session_state.get(key_grade, default_grade)
                cgrade_col.markdown(f"<div class='small-selected'>{display_grade}</div>", unsafe_allow_html=True)
                grade_options = [""] + grade_set
                try:
                    target = st.session_state.get(key_grade, default_grade) or ""
                    matched_index = 0
                    for idx_opt, opt in enumerate([""] + grade_set):
                        if not opt:
                            continue
                        if normalize(opt) == normalize(str(target)):
                            matched_index = idx_opt
                            break
                    grade_sel = cgrade_col.selectbox("", options=grade_options, index=matched_index, key=key_grade, help="Grade (editable)")
                except Exception:
                    grade_sel = cgrade_col.selectbox("", options=grade_options, index=0, key=key_grade, help="Grade (editable)")
                df.at[i, f"Grade_{t_idx+1}"] = st.session_state.get(key_grade, grade_sel) or ""

                # Finish
                fin_set = []
                for catk, items in din_index.items():
                    for it in items:
                        if it.get("type_name") != chosen_type:
                            continue
                        
                        for f in it.get("finishes", []):
                            if f:
                                fup = str(f).strip().upper()
                                if fup not in fin_set:
                                    fin_set.append(fup)
                for f in finish_suggestions[i]:
                    fu = str(f).strip().upper()
                    if fu and fu not in fin_set:
                        fin_set.append(fu)
                for f in din_global_finishes:
                    if f and f not in fin_set:
                        fin_set.append(f)
                default_fin = df.at[i, f"Finish_{t_idx+1}"] or ""
                try:
                    for fin in fin_set:
                        if normalized_token_in_text(fin, desc_for_match):
                            default_fin = fin
                            df.at[i, f"Finish_{t_idx+1}"] = fin
                            st.session_state[key_finish] = fin
                            break
                except Exception:
                    pass
                display_fin = st.session_state.get(key_finish, default_fin)
                cfinish_col.markdown(f"<div class='small-selected'>{display_fin}</div>", unsafe_allow_html=True)
                fin_options_full = [""] + fin_set
                try:
                    fin_sel = cfinish_col.selectbox("", options=fin_options_full, index=fin_options_full.index(default_fin) if default_fin in fin_options_full else 0, key=key_finish, help="Finish (editable)")
                except Exception:
                    fin_sel = cfinish_col.selectbox("", options=fin_options_full, index=0, key=key_finish, help="Finish (editable)")
                df.at[i, f"Finish_{t_idx+1}"] = st.session_state.get(key_finish, fin_sel) or ""

                first_block = False

            chosen_types_for_row = []
            for idx in range(len(detected_cats)):
                tval = df.at[i, f"Type_{idx+1}"] if f"Type_{idx+1}" in df.columns else ""
                if tval:
                    chosen_types_for_row.append(tval)
            df.at[i, "Type"] = ",".join(chosen_types_for_row)
            cat_str = ",".join(detected_cats) if detected_cats else ""
            stds = []
            for idx in range(len(detected_cats)):
                sval = df.at[i, f"Standard_{idx+1}"] if f"Standard_{idx+1}" in df.columns else ""
                if sval:
                    stds.append(sval)
            df.at[i, "DIN Details"] = f"{cat_str}/{df.at[i,'Type']}/{','.join(stds)}"

        st.markdown("---")
        st.markdown("### Finalize and generate merged table")
        if st.button("Generate Final Table"):
            final_rows = []
            for i, row in df.iterrows():
                desc_text = str(row[desc_col])
                std_parts = []
                grade_parts = []
                finish_parts = []
                type_cols = [c for c in df.columns if re.match(r"Type_\d+$", c)]
                type_cols_sorted = sorted(type_cols, key=lambda x: int(x.split("_")[1]))  # Type_1, Type_2 ...
                for idx_col, tcol in enumerate(type_cols_sorted, start=1):
                    sval = row.get(f"Standard_{idx_col}", "")
                    if sval and sval not in std_parts:
                        std_parts.append(str(sval))
                    gval = row.get(f"Grade_{idx_col}", "")
                    if gval and gval not in grade_parts:
                        grade_parts.append(str(gval))
                    fval = row.get(f"Finish_{idx_col}", "")
                    if fval and fval not in finish_parts:
                        finish_parts.append(str(fval))
                if not std_parts and row.get("Standard"):
                    std_parts.append(row.get("Standard"))
                if not grade_parts and row.get("Grade"):
                    grade_parts.append(row.get("Grade"))
                if not finish_parts and row.get("Finish"):
                    finish_parts.append(row.get("Finish"))

                std_merged = " / ".join(std_parts) if std_parts else ""
                grade_merged = " / ".join(grade_parts) if grade_parts else ""
                finish_merged = " / ".join(finish_parts) if finish_parts else ""

                final_rows.append({
                    "Sr No": i+1,
                    "Description": desc_text,
                    "Dimension Standard": std_merged,
                    "Grade": grade_merged,
                    "Finish": finish_merged
                })

            final_df = pd.DataFrame(final_rows, columns=["Sr No", "Description", "Dimension Standard", "Grade", "Finish"])
            st.markdown("### Final Table")
            st.dataframe(final_df, use_container_width=True)

            # --- ✅ Automatically pass final table to Tab 5 ---
            st.session_state["mapped_df"] = final_df.copy()
            st.caption("✅ Final mapped table automatically sent to Tab 5 (Rate / Weight Calculator)")

        else:
            st.info("When you're ready, click 'Generate Final Table' to merge per-type columns into the final 5-column table.")

    
# ----------------------------
# TAB 4: NOTES & TERMS SECTION
# ----------------------------
with tab4:
    st.subheader("📝 Notes")
    st.session_state["additional_note"] = st.text_area("Additional Note", "")

    # ----------------------------
    # FIXED MAIN NOTE (Always on top)
    # ----------------------------
    MAIN_NOTE = (
        "PRICES ARE BASED SOLELY ON THE STATED QUANTITIES.ANY ALTERATION IN QUANTITIES MAY LEAD TO A CHANGE IN PRICING.A REVISED QUOTATION WILL BE ISSUED AS NEEDED."
    )
    st.markdown(f"**{MAIN_NOTE}**")
    st.markdown("---")

    # ----------------------------
    # SUB NOTES (Checkbox)
    # ----------------------------
    st.markdown("### Additional Notes (Select Any)")

    SUB_NOTES = [
        "1) 10% + or - in qty shall be accepted.",
        "2) Kindly find highlighted changes in red colour if any before confirming the order.",
        "3) Material Test Certificate will be provided as per EN 10204 3.1."
    ]

    selected_subnotes = []
    for note in SUB_NOTES:
        if st.checkbox(note, key=f"sub_{note[:10]}"):
            selected_subnotes.append(note)

    st.markdown("---")

    # ----------------------------
    # PART A — TERMS & CONDITIONS
    # ----------------------------
    st.markdown("### Part A — Terms & Conditions")

    # ---- DEFAULT OPTIONS (always present in dropdown) ----
    DELIVERY_DEFAULT = [
        "EX-MUMBAI,CUSTOMER TRANSPORT",
        "EX WORKS",
        "DELIVERY INCLUSIVE",
        "GO DOWN DELIVERY"
    ]

    PAYMENT_DEFAULT = [
        "IMMEDIATE",
        "30% ADVANCE BALANCE AGAINST PI",
        "30 DAYS CREDIT",
        "45 DAYS CREDIT",
        "100% ADVANCE AGAINST PI"
    ]

    VALIDITY_DEFAULT = [
        "1 DAY",
        "1 WEEK",
        "1 MONTH"
    ]

    # ---- Ensure session_state sets exist ----
    if "delivery_options" not in st.session_state:
        st.session_state.delivery_options = set([opt.upper() for opt in DELIVERY_DEFAULT])

    if "payment_options" not in st.session_state:
        st.session_state.payment_options = set([opt.upper() for opt in PAYMENT_DEFAULT])

    if "validity_options" not in st.session_state:
        st.session_state.validity_options = set([opt.upper() for opt in VALIDITY_DEFAULT])

    colA1, colA2 = st.columns(2)

    # ----------------------------
    # LEFT COLUMN (Delivery + Payment)
    # ----------------------------
    with colA1:

        # ---- DELIVERY ----
        delivery_choice = st.selectbox(
            "DELIVERY",
            sorted(st.session_state.delivery_options) + ["➕ ADD NEW"]
        )
        if delivery_choice == "➕ ADD NEW":
            new_delivery = st.text_input("Enter new Delivery term")
            if new_delivery:
                new_delivery = new_delivery.upper()
                st.session_state.delivery_options.add(new_delivery)
                delivery_choice = new_delivery

        # ---- PAYMENT ----
        payment_choice = st.selectbox(
            "PAYMENT",
            sorted(st.session_state.payment_options) + ["➕ ADD NEW"]
        )
        if payment_choice == "➕ ADD NEW":
            new_payment = st.text_input("Enter new Payment term")
            if new_payment:
                new_payment = new_payment.upper()
                st.session_state.payment_options.add(new_payment)
                payment_choice = new_payment

    # ----------------------------
    # RIGHT COLUMN (Validity + Fixed Values)
    # ----------------------------
    with colA2:

        period = "4-5 WEEKS"
        tax = " GST 18%"
        pf = "3%"

        # ---- VALIDITY ----
        validity_choice = st.selectbox(
            "VALIDITY",
            sorted(st.session_state.validity_options) + ["➕ ADD NEW"]
        )
        if validity_choice == "➕ ADD NEW":
            new_validity = st.text_input("Enter new Validity")
            if new_validity:
                new_validity = new_validity.upper()
                st.session_state.validity_options.add(new_validity)
                validity_choice = new_validity

    # Showing fixed terms
    st.markdown(f"**{period}**")
    st.markdown(f"**{tax}**")
    st.markdown(f"**{pf}**")

    st.markdown("---")
    # ----------------------------
    # PART B — BANK DETAILS (Fixed)
    # ----------------------------
    st.markdown("### Part B — Bank Details (Fixed)")

    BANK_DETAILS = [
        "BANK NAME : KOTAK MAHINDRA BANK.",
        "ADD : GROUND AND MEZZANINE FLOOR,",
        "BOTAWALA CHAMBER,2, MUMBAI- 01.",
        "BRANCH CODE - 957",
        "IFSC CODE NO KKBK0000957",
        "OUR A/C NO 9223312803."
    ]


    for line in BANK_DETAILS:
        st.write(line)

    st.markdown("---")

    # ----------------------------
    # FOOTER LINES (Fixed)
    # ----------------------------
    FOOTER_LINES = [
    "Thank you for reviewing our quotation.",
    "For any queries or clarification, our team is always available to assist you.",
    "We look forward to your positive consideration and hope to move ahead together.",
    ]


    st.markdown("### Footer")
    for line in FOOTER_LINES:
        st.write(line)

    # ----------------------------
    # STORE EVERYTHING IN SESSION
    # ----------------------------
    st.session_state["quotation_notes"] = {
        "main_note": MAIN_NOTE,
        "sub_notes": selected_subnotes,
        "delivery": delivery_choice,
        "period": period,
        "tax": tax,
        "pf": pf,
        "payment": payment_choice,
        "validity": validity_choice,
        "bank_details": BANK_DETAILS,
        "footer": FOOTER_LINES
    }


# ----------------------------
# TAB 5: RATE / WEIGHT CALCULATOR
# ----------------------------
if "df" in locals():
    with tab5:
        st.title("🔩 Rate / Weight Calculator")
        if "last_message" in st.session_state:
            st.info(st.session_state["last_message"])
            del st.session_state["last_message"]

        # --- ✅ Auto-load mapped table from Tab 3 if available ---
        if "mapped_df" in st.session_state:
            df = st.session_state["mapped_df"].copy()
            st.success("✅ Loaded mapped table from Tab 3.")

            # --- ✅ Normalize column names for compatibility ---
            rename_map = {
                "Description": "Material Description",
                "Sr No": "Sr. No.",
                "Dimension Standard": "Dimension\nStandard",
                "Grade": "Material\nGrade"
            }
            df.rename(columns=rename_map, inplace=True)
        else:
            st.warning("⚠️ No mapped table found. Please complete mapping in Tab 3 first.")
            st.stop()

        # -------------------------
        # ✅ Keep mapped table from Tab 3 persistent and editable
        # -------------------------
        if "mapped_df" in st.session_state:
           # Only initialize once
           if "items_df" not in st.session_state:
              mapped = st.session_state["mapped_df"].copy()
              mapped.rename(columns=rename_map, inplace=True)
          
              if "Rate" not in mapped.columns:
               mapped["Rate"] = ""
               st.session_state["items_df"] = mapped
        else:
    
            if "items_df" not in st.session_state:
               st.session_state["items_df"] = pd.DataFrame(columns=["Material Description", "Rate"])


        edited_df = st.session_state["items_df"]
        # -----------------------------------------
        # 🔽 DISCOUNT SECTION (GLOBAL DISCOUNT %)
        # -----------------------------------------
        st.markdown("### 💸 Discount")

# ----- Row layout -----
        col1, col2, col3, col4 = st.columns([2, 2, 2, 2])

# ---- Discount Input ----
        with col1:
            discount_percent = st.number_input(
            "Discount (%)",
            min_value=0.0, max_value=100.0,
            step=0.1,
            value=st.session_state.get("global_discount", 0.0),
            key="discount_percent"
        )
        st.session_state["global_discount"] = discount_percent

# ---- Apply All button (inside container → no lag) ----
        with col2:
            with st.container():
                if st.button("Apply All", key="apply_all_btn"): 
                    st.session_state["discount_mode"] = "global"
                    st.session_state["global_discount"] = discount_percent

                    df = st.session_state["items_df"].copy()
                    for idx in df.index:
                        try:
                            r = float(df.at[idx, "Rate"])
                            df.at[idx, "Rate"] = f"{r - (r * discount_percent / 100):.2f}"
                        except:
                             pass
                    st.session_state["items_df"] = df
                    st.success("✔ Applied to all rows")
                    st.rerun()

        # ---- Row No ----
        with col3:
            row_to_apply = st.number_input(
            "Row No",
            min_value=1,
            max_value=len(st.session_state["items_df"]),
            step=1,
            key="discount_row"
        )

        # ---- Apply Row button (container) ----
        with col4:
            with st.container():
                if st.button("Apply Row", key="apply_row_btn"):
                    st.session_state["discount_mode"] = "row"  
                    idx = row_to_apply - 1
                    df = st.session_state["items_df"].copy()
                    try:
                        r = float(df.at[idx, "Rate"])
                        df.at[idx, "Rate"] = f"{r - (r * discount_percent / 100):.2f}"
                        st.session_state["items_df"] = df
                        st.success(f"✔ Discount applied to row {row_to_apply}")
                    except:
                        st.error("❌ Invalid row rate!")
                    st.rerun()

        # === Top: Editable table pinned at the top, compact ===
        st.markdown('<div class="tiny">Finish goods description (editable table below)</div>', unsafe_allow_html=True)
        if "Rate" not in edited_df.columns:
            edited_df["Rate"] = ""
        df_calc = st.session_state["items_df"].copy()

        def to_float(x):
            try:
                return float(str(x).strip())
            except:
                return 0.0

        df_calc["Rate_float"] = df_calc["Rate"].apply(to_float)

        mode = st.session_state.get("discount_mode", "none")
        discount_percent = st.session_state.get("global_discount", 0.0)
        if mode == "global":
    # Reverse = discounted / (1 - p/100)
           df_calc["Original_before_discount"] = df_calc["Rate_float"] / (1 - (discount_percent / 100))
           subtotal = df_calc["Original_before_discount"].sum()
        else:
           subtotal = None  # subtotal hidden for row mode
        grand_total = df_calc["Rate_float"].sum()
        if mode == "global":
            discount_amount = subtotal - grand_total
        else:
            discount_amount = None
        
        show_summary = st.checkbox("Show Totals (Subtotal / Discount / Grand Total)", 
                           value=True, 
                           key="show_summary_checkbox")
        if show_summary:
            st.subheader("📊 Summary")

    # CASE 1 – Global discount (Apply All)
            if mode == "global":
                st.write(f"**Subtotal (Before Discount):** ₹ {subtotal:,.2f}")
                st.write(f"**Discount ({discount_percent}%):** ₹ {discount_amount:,.2f}")
                st.write(f"**Grand Total:** ₹ {grand_total:,.2f}")

    # CASE 2 – Row discounts or no discounts
            else:
                st.write(f"**Grand Total:** ₹ {grand_total:,.2f}")


        # ✅ Directly edit session DataFrame (live sync)
        st.session_state["items_df"] = st.data_editor(
            st.session_state["items_df"],
            num_rows="dynamic",
            use_container_width=True
        )
        edited_df = st.session_state["items_df"]

        st.markdown("---")

        # Compact CSS to reduce wasted space and provide subtle panels
        st.markdown(
            """
            <style>
            .panel {background: #f0f7fb; border-radius: 6px; padding: 6px 10px; margin-bottom: 6px; font-size:13px}
            .panel-soft {background: #f7fbf0; border-radius: 6px; padding: 6px 10px; margin-bottom: 6px; font-size:13px}
            .muted {color: #666; font-size: 12px;}
            .small-note {font-size:12px; color:#666; margin-top:2px}
            .tiny {font-size:12px; color:#444; margin-bottom:6px}
            .compact-info {font-size:13px; color:#444; margin-bottom:4px}
            .no-pad {padding:0; margin:0}
            .data-editor .stTextInput>div>input {height:36px}
            </style>
            """,
            unsafe_allow_html=True,
        )

        # -------------------------
        # Auto-load JSONs from project folder
        # -------------------------
        import os, json
        base_dir = os.path.dirname(os.path.abspath(__file__))
        catalogue_path = os.path.join(base_dir, "catalogue.json")
        gross_path = os.path.join(base_dir, "gross_weight_data.json")

        catalogue = None
        dia_data = None
        if os.path.exists(catalogue_path):
            with open(catalogue_path, "r", encoding="utf-8") as f:
                try:
                    catalogue = load_catalogue(f)
                except Exception:
                    f.seek(0)
                    catalogue = json.load(f)
        else:
            st.error("❌ catalogue.json not found next to weight.py — place it there.")

        if os.path.exists(gross_path):
            with open(gross_path, "r", encoding="utf-8") as f:
                dia_data = json.load(f)
        else:
            st.error("❌ gross_weight_data.json not found next to weight.py — place it there.")

        # -------------------------
        # Keywords for combo detection (as requested)
        # -------------------------
        PRIMARY_KEYWORDS = ["STUD", "NUT", "BOLT", "WASHER", "SCREW"]

        # -------------------------
        # Helper: normalize string for matching
        # -------------------------
        def _norm(s):
            if s is None:
                return ""
            return str(s).strip().upper()

        # -------------------------
        # Helper: get all screw_type labels from catalogue that belong to a primary category
        # -------------------------
        def get_types_for_category(category, catalogue_list):
            cat_l = category.lower()
            labels = []
            if cat_l == "washer":
                for entry in catalogue_list:
                    stype = entry.get("screw_type", "")
                    if stype and "washer" in stype.lower():
                        labels.append(stype)
            else:
                for entry in catalogue_list:
                    stype = entry.get("screw_type", "")
                    if not stype:
                        continue
                    if cat_l in stype.lower():
                        labels.append(stype)
                if not labels:
                    for entry in catalogue_list:
                        stype = entry.get("screw_type", "")
                        if not stype:
                            continue
                        tokens = [t.strip().lower() for t in stype.replace("-", " ").replace("_", " ").split()]
                        if cat_l in tokens:
                            labels.append(stype)
            seen = set()
            out = []
            for l in labels:
                if l not in seen:
                    out.append(l)
                    seen.add(l)
            return out

        # -------------------------
        # Helper: Detect primary categories present in description (keyword-based)
        # -------------------------
        def detect_primary_categories(description, keywords=PRIMARY_KEYWORDS):
            desc = _norm(description)
            found = []
            for kw in keywords:
                if kw in desc:
                    found.append(kw)
            token_map = {
                "CAPSCREW": "SCREW",
                "CAP SCREW": "SCREW",
                "ALLEN": "SCREW",
                "ALLENCAP": "SCREW",
                "ALLEN CAP": "SCREW",
                "HEX HD": "BOLT",
                "HEX HEAD": "BOLT",
                "HEX": "BOLT",
                "FLAT WASHER": "WASHER",
                "SPRING WASHER": "WASHER",
            }
            for token, mapped in token_map.items():
                if token in desc and mapped not in found:
                    found.append(mapped)
            if "WASHER" in desc and "WASHER" not in found:
                found.append("WASHER")
            return found

        # -------------------------
        # Smart primary-type matching fallback using rapidfuzz if needed
        # -------------------------
        try:
            from rapidfuzz import process, fuzz
        except Exception:
            process = None
            fuzz = None

        def smart_match_primary_type(description: str, screw_types: list, threshold: float = 70.0):
            desc = description or ""
            desc_l = desc.lower()
            for idx, s in enumerate(screw_types):
                if s and s.lower() in desc_l:
                    return [(s, 100.0, idx)]
            try:
                if process is None:
                    # fallback simple substring matching
                    out = []
                    for idx, s in enumerate(screw_types):
                        sc = 100.0 if (s and s.lower() in desc_l) else 0.0
                        out.append((s, sc, idx))
                    out_sorted = sorted(out, key=lambda x: x[1], reverse=True)
                    return out_sorted[:3]
                raw_matches = process.extract(description, screw_types, scorer=fuzz.WRatio, limit=6)
                good = [(m, float(sc), idx) for m, sc, idx in raw_matches if sc >= threshold]
                if good:
                    return good
                return [(m, float(sc), idx) for m, sc, idx in raw_matches[:3]]
            except Exception:
                return []

        # -------------------------
        # Helper: safely update Rate + set next unfilled item
        # -------------------------
        def update_rate_and_next(selected_desc, value):
            try:
                from rapidfuzz import process, fuzz

                df = st.session_state["items_df"]

        # list of all descriptions in table
                all_descs = df["Material Description"].astype(str).tolist()

        # 🔍 fuzzy match to find closest row (score >= 70)
                match = process.extractOne(
                    selected_desc,
                    all_descs,
                    scorer=fuzz.WRatio
                )

                if match:
                    matched_desc, score, idx = match

            # Only update if confidence is reasonable
                    if score >= 70:
                        df.at[idx, "Rate"] = value
                        st.session_state["items_df"] = df
                        st.session_state["last_message"] = (
                            f"✅ Mapped {value} to '{matched_desc}' (match score {score})"
                        )
                    else:
                        st.session_state["last_message"] = (
                            f"⚠️ No strong match found for '{selected_desc}'. (Score {score})"
                        )
                else:
                    st.session_state["last_message"] = (
                        f"❌ No match found for '{selected_desc}'."
                    )

        # ---- Find next item with blank rate ----
                remaining = [
                    d for d in df["Material Description"]
                    if pd.notna(d)
                        and str(d).strip() != ""
                        and (str(df.loc[df["Material Description"] == d, "Rate"].iloc[0]).strip() == "")
                ]
                st.session_state["next_item"] = remaining[0] if remaining else None
                st.rerun()

            except Exception as e:
                st.session_state["last_message"] = f"⚠️ Update failed: {e}"
                st.rerun()

        # -------------------------
        # Mode selector (compact dropdown)
        # -------------------------
        weight_mode = st.selectbox("Select Mode", ["Finish Goods Weight", "Gross Weight"])


        # -------------------------
        # Finish Goods Mode
        # -------------------------
        if weight_mode == "Finish Goods Weight":
            st.markdown('<div class="panel"><b class="tiny">Finish Goods Mode</b> — edit items & map computed rates/weights. Use <span class="tiny">➕ Add Value</span> to write Rate and jump to next unfilled item.</div>', unsafe_allow_html=True)

            if not catalogue:
                st.info("No catalogue loaded. Place catalogue.json beside weight.py.")
                st.stop()

            if "next_item" in st.session_state and pd.notna(st.session_state["next_item"]):
                selected_desc = st.session_state["next_item"]
            else:
                selected_desc = None
                for desc in edited_df["Material Description"]:
                    if pd.notna(desc) and desc != "":
                        selected_desc = desc
                        break

            if not selected_desc:
                st.info("No Material Description in table. Add items and retry.")
                st.stop()

            top1, top2 = st.columns([8,4])
            with top1:
                selected_desc = st.text_input("Current Material Description (Editable)", value=selected_desc)
            with top2:
                st.markdown('<div class="small-note">Tip: edit the description to refine auto-matching.</div>', unsafe_allow_html=True)

            st.markdown("---")

            screw_input_to_use = selected_desc
            combo_items = parse_combo_input(screw_input_to_use)

            if not combo_items or len(combo_items) <= 1:
                detected = detect_primary_categories(screw_input_to_use, PRIMARY_KEYWORDS)
                if len(detected) > 1:
                    combo_items = [{"type": cat.lower(), "count": 1} for cat in detected]

            # -------------------------
            # Combo Mode
            # -------------------------
            if len(combo_items) > 1:
                st.markdown('<div class="panel-soft"><b class="tiny">⚙️ Combo Mode</b></div>', unsafe_allow_html=True)
                total_weight_pc, total_price_all = 0.0, 0.0

                for i, comp in enumerate(combo_items):
                    ctype = comp.get("type", "")
                    header_label = str(ctype).title()
                    st.markdown(f"### 🔹 {header_label} (x{comp.get('count',1)})")

                    type_options = get_types_for_category(ctype, catalogue)
                    if not type_options:
                        all_labels = [e.get("screw_type", "") for e in catalogue if e.get("screw_type")]
                        fuzzy_matches = smart_match_primary_type(ctype, all_labels, threshold=40.0)
                        type_options = [m for m, sc, idx in fuzzy_matches]

                    if not type_options:
                        st.warning(f"No catalogue options found for {header_label}")
                        continue

                    chosen_type = st.selectbox(f"Select {header_label} Type", type_options, key=f"combo_type_{i}")
                    entries = [e for e in catalogue if e.get("screw_type") == chosen_type]

                    if not entries:
                        st.warning(f"No catalogue entries for selected type '{chosen_type}'")
                        continue

                    sel_c1, sel_c2, sel_c3 = st.columns([2,3,3])
                    with sel_c1:
                        has_metric = any(e.get("dimensions_in_metric") for e in entries)
                        has_inches = any(e.get("dimensions_in_inches") for e in entries)
                        dim_type = "metric"
                        if has_metric and has_inches:
                            dim_type = st.radio("", ["metric", "inches"], index=0, key=f"combo_dim_{i}")
                        elif has_inches:
                            dim_type = "inches"
                    with sel_c2:
                        diam_list = []
                        for e in entries:
                            dims = e.get("dimensions_in_metric", []) if dim_type == "metric" else e.get("dimensions_in_inches", [])
                            for d in dims:
                                diam_list.extend(list(d.get("diameter", {}).keys()))
                        diam_list = sort_diameters(diam_list, dim_type)
                        diameter = st.selectbox("Diameter", diam_list, key=f"combo_dia_{i}")
                    with sel_c3:
                        lengths = []
                        for e in entries:
                            dims = e.get("dimensions_in_metric", []) if dim_type == "metric" else e.get("dimensions_in_inches", [])
                            for d in dims:
                                if diameter in d.get("diameter", {}):
                                    length_val = d.get("length_mm") or d.get("length")
                                    if length_val is not None:
                                        lengths.append(length_val)
                        lengths = sorted(set(lengths), key=lambda x: int(x) if isinstance(x, (int, str)) and str(x).isdigit() else str(x))
                        if lengths:
                            if len(lengths) == 1:
                                length_val = lengths[0]
                                st.info(f"Auto length: {length_val}")
                            else:
                                length_choice = st.selectbox("Length", ["--manual--"] + [str(l) for l in lengths], key=f"combo_len_{i}")
                                if length_choice == "--manual--":
                                    if dim_type == "metric":
                                        length_val = st.number_input("Enter length (mm)", min_value=1, step=1, key=f"combo_manual_len_{i}")
                                    else:
                                        length_val = st.text_input("Enter length (inches)", key=f"combo_manual_len_{i}")
                                else:
                                    length_val = int(length_choice) if dim_type == "metric" and str(length_choice).isdigit() else length_choice
                        else:
                            length_val = None
                            st.info("No length for this diameter.")

                    r1, r2, r3 = st.columns([3,2,2])
                    with r1:
                        rate_val = find_rate(entries, length_val, diameter, dim_type)
                        if not rate_val:
                            st.warning("No rate found for this selection.")
                            continue
                    with r2:
                        unit = entries[0].get("unit", "").lower()
                        if "price per piece" in unit:
                            per_pc_wt = float(rate_val)
                            st.success(f"Per pc: {per_pc_wt:.4f} kg")
                        elif "approx. weight per 100" in unit or "100 nos" in unit:
                            per_pc_wt = float(rate_val) / 100
                            st.success(f"Per 100: {float(rate_val)} kg")
                        elif "approx. count per 50" in unit:
                            pcs_per_50kg = float(rate_val)
                            per_pc_wt = 50.0 / pcs_per_50kg
                            st.success(f"Per 50kg: {int(pcs_per_50kg)} pcs")
                        else:
                            per_pc_wt = float(rate_val)
                            st.info("Raw used.")
                    with r3:
                        rate_per_kg = st.number_input("Rate/kg", min_value=0.0, step=0.01, key=f"combo_rate_{i}")
                        qty = st.number_input("Qty", min_value=1, value=1, key=f"combo_qty_{i}")

                    comp_weight = per_pc_wt * comp.get("count", 1) * qty
                    comp_price = comp_weight * rate_per_kg
                    total_weight_pc += per_pc_wt * comp.get("count", 1)
                    total_price_all += comp_price

                    s1, s2, s3 = st.columns(3)
                    s1.success(f"{header_label} wt/pc: {per_pc_wt:.4f} kg")
                    s2.success(f"{header_label} total wt: {comp_weight:.4f} kg")
                    s3.success(f"{header_label} total ₹: {comp_price:.2f}")

                st.markdown("---")
                st.success(f"⚖️ Combined per-piece weight: {total_weight_pc:.4f} kg")
                st.success(f"💰 Combined total price: ₹{total_price_all:.2f}")

                # Use update_rate_and_next to write the total price into Rate for the selected_desc
                if st.button("➕ Add Value"):
                    total_price_str = f"{total_price_all:.2f}"
                    update_rate_and_next(selected_desc, total_price_str)


            # -------------------------
            # Single Item Mode
            # -------------------------
            else:
                st.markdown('<div class="panel"><span class="tiny">🔧 Single Item Mode</span></div>', unsafe_allow_html=True)

                screw_types = [item.get("screw_type", "") for item in catalogue if "screw_type" in item]

                matches = smart_match_primary_type(screw_input_to_use or "", screw_types, threshold=70.0)

                if not matches:
                    st.error("No good match. Try clearer description.")
                    st.stop()

                match_labels = [f"{m} ({sc:.1f}%)" for m, sc, _ in matches]
                col_a, col_b = st.columns([6,6])
                with col_a:
                    chosen_match = st.selectbox("Closest Type", match_labels, key="closest_type_single")
                with col_b:
                    best_match_name = matches[match_labels.index(chosen_match)][0]
                    best_match_base = normalize_name(best_match_name)
                    matched_entries = [item for item in catalogue if normalize_name(item.get("screw_type", "")) == best_match_base]
                    entry_labels = [f"{entry.get('screw_type','Unknown')} | Standard: {entry.get('standard','N/A')} | Unit: {entry.get('unit','N/A')}" for entry in matched_entries]
                    if entry_labels:
                        chosen_entry_label = st.selectbox("Select Entry (Standard)", entry_labels, key="entry_select_single")
                        chosen_entry_idx = entry_labels.index(chosen_entry_label)
                        relevant_entries = [matched_entries[chosen_entry_idx]]
                    else:
                        st.warning("No entries available for chosen type.")
                        relevant_entries = []

                has_metric = any(entry.get("dimensions_in_metric") and len(entry.get("dimensions_in_metric")) > 0 for entry in relevant_entries)
                has_inches = any(entry.get("dimensions_in_inches") and len(entry.get("dimensions_in_inches")) > 0 for entry in relevant_entries)

                if has_metric and has_inches:
                    dim_type = st.radio("", ["metric", "inches"], key="dim_type_single")
                elif has_metric:
                    dim_type = "metric"
                    st.markdown('<div class="compact-info">Only metric available.</div>', unsafe_allow_html=True)
                elif has_inches:
                    dim_type = "inches"
                    st.markdown('<div class="compact-info">Only inches available.</div>', unsafe_allow_html=True)
                else:
                    st.error("No dimension data.")
                    st.stop()

                diameters = []
                if dim_type == "metric":
                    for entry in relevant_entries:
                        for dim in entry.get("dimensions_in_metric", []):
                            diameters.extend(list(dim.get("diameter", {}).keys()))
                else:
                    for entry in relevant_entries:
                        for dim in entry.get("dimensions_in_inches", []):
                            diameters.extend(list(dim.get("diameter", {}).keys()))
                diameters = sort_diameters(diameters, dim_type)

                diameter_choice = st.selectbox("Diameter", ["--manual--"] + diameters, key="diam_choice_single")
                if diameter_choice == "--manual--":
                    diameter = st.text_input("Enter Diameter (e.g. M6)", key="diam_manual_single")
                else:
                    diameter = diameter_choice

                lengths = []
                has_length_values = False
                if dim_type == "metric":
                    for entry in relevant_entries:
                        for dim in entry.get("dimensions_in_metric", []):
                            if diameter in dim.get("diameter", {}):
                                if dim.get("length_mm") is not None:
                                    lengths.append(dim.get("length_mm"))
                                    has_length_values = True
                else:
                    for entry in relevant_entries:
                        for dim in entry.get("dimensions_in_inches", []):
                            if diameter in dim.get("diameter", {}):
                                if dim.get("length") is not None:
                                    lengths.append(dim.get("length"))
                                    has_length_values = True
                lengths = sorted(set(lengths), key=lambda x: int(x) if isinstance(x, (int, str)) and str(x).isdigit() else str(x))

                if not has_length_values:
                    dcol1, dcol2, dcol3 = st.columns([4,4,4])
                    with dcol2:
                        rate_per_kg = st.number_input("Rate/kg", min_value=0.0, step=0.01, format="%.2f", key="rate_per_kg_single_no_len")
                    with dcol3:
                        quantity = st.number_input("Qty", min_value=1, value=1, key="qty_single_no_len")
                    rate_value = find_rate(relevant_entries, None, diameter, dim_type) if relevant_entries else None
                else:
                    if len(lengths) == 1:
                        length_val = lengths[0]
                        st.info(f"Auto-length: {length_val}")
                    else:
                        length_choice = st.selectbox("Length", ["--manual--"] + [str(l) for l in lengths], key="length_choice_single")
                        if length_choice == "--manual--":
                            if dim_type == "metric":
                                length_val = st.number_input("Enter length (mm)", min_value=1, step=1, key="manual_length_single")
                            else:
                                length_val = st.text_input("Enter length (inches)", key="manual_length_single")
                        else:
                            length_val = int(length_choice) if dim_type == "metric" and str(length_choice).isdigit() else length_choice

                    wcol1, wcol2, wcol3 = st.columns([4,4,4])
                    with wcol1:
                        rate_value = find_rate(relevant_entries, length_val, diameter, dim_type)
                    with wcol2:
                        rate_per_kg = st.number_input("Rate/kg", min_value=0.0, step=0.01, format="%.2f", key="rate_per_kg_single")
                    with wcol3:
                        quantity = st.number_input("Qty", min_value=1, value=1, key="qty_single")

                btn_col1, btn_col2 = st.columns([3,3])
                with btn_col1:
                    calculate_pressed = st.button("Calculate", key="calc_single")
                with btn_col2:
                    add_value_pressed = st.button("➕ Add Value", key="add_single")

                if calculate_pressed:
                    if 'rate_value' in locals() and rate_value:
                        chosen_unit = relevant_entries[0].get("unit", "").strip().lower() if relevant_entries else ""
                        if "price per piece" in chosen_unit:
                            weight_per_pc = float(rate_value)
                            st.success(f"Per Piece Weight: {weight_per_pc:.4f} kg")
                        elif "approx. count per 50" in chosen_unit:
                            pcs_per_50kg = float(rate_value)
                            weight_per_pc = 50.0 / pcs_per_50kg
                            st.success(f"Pieces per 50kg: {int(pcs_per_50kg)}")
                        elif "approx. weight per 100" in chosen_unit or "100 nos" in chosen_unit:
                            weight_per_100 = float(rate_value)
                            weight_per_pc = weight_per_100 / 100.0
                            st.success(f"Weight per 100: {weight_per_100} kg")
                        else:
                            try:
                                pcs_per_50kg = float(rate_value)
                                weight_per_pc = 50.0 / pcs_per_50kg
                                st.success(f"Pieces per 50kg (inferred): {int(pcs_per_50kg)}")
                            except Exception:
                                weight_per_pc = float(rate_value)
                                st.success(f"Weight/pc (inferred): {weight_per_pc:.4f} kg")

                        total_weight = weight_per_pc * quantity
                        per_piece_price = weight_per_pc * (locals().get("rate_per_kg", locals().get("rate_per_kg_single_no_len", 0.0)))
                        total_price = per_piece_price * quantity

                        sum_c1, sum_c2, sum_c3 = st.columns(3)
                        sum_c1.success(f"Weight/pc: {weight_per_pc:.4f} kg")
                        sum_c2.success(f"Total wt: {total_weight:.4f} kg")
                        sum_c3.success(f"Total ₹: {total_price:.2f}")
                    else:
                        st.error("No matching rate found for this diameter/length.")

                if add_value_pressed:
                    total_to_write = None
                    try:
                        if 'rate_value' in locals() and rate_value:
                            chosen_unit = relevant_entries[0].get("unit", "").strip().lower() if relevant_entries else ""
                            if "price per piece" in chosen_unit:
                                weight_per_pc = float(rate_value)
                            elif "approx. count per 50" in chosen_unit:
                                pcs_per_50kg = float(rate_value)
                                weight_per_pc = 50.0 / pcs_per_50kg
                            elif "approx. weight per 100" in chosen_unit or "100 nos" in chosen_unit:
                                weight_per_100 = float(rate_value)
                                weight_per_pc = weight_per_100 / 100.0
                            else:
                                try:
                                    pcs_per_50kg = float(rate_value)
                                    weight_per_pc = 50.0 / pcs_per_50kg
                                except Exception:
                                    weight_per_pc = float(rate_value)
                            rp = locals().get("rate_per_kg", locals().get("rate_per_kg_single_no_len", 0.0))
                            per_piece_price = weight_per_pc * rp
                            total_price = per_piece_price * locals().get("quantity", 1)
                            total_to_write = f"{total_price:.2f}"
                        else:
                            rp = locals().get("rate_per_kg", locals().get("rate_per_kg_single_no_len", None))
                            if rp is not None:
                                total_to_write = f"{float(rp) * locals().get('quantity',1):.2f}"
                            else:
                                total_to_write = "0.00"
                    except Exception:
                        total_to_write = "0.00"

                    # Use update_rate_and_next to write and advance
                    update_rate_and_next(selected_desc, total_to_write)


        # -------------------------
        # Gross Weight Mode (NEW JSON STRUCTURE)
        # -------------------------
        elif weight_mode == "Gross Weight":
            st.markdown('<div class="panel"><b class="tiny">⚖️ Gross Weight Calculator</b> — Single formula mode.</div>', unsafe_allow_html=True)

            if not dia_data:
                st.warning("gross_weight_data.json missing. Place it next to weight.py.")
                st.stop()

    # Convert JSON list → DataFrame for easy filtering
            df_gross = pd.DataFrame(dia_data)

    # --- UI ---
            g1, g2, g3 = st.columns([3,2,3])
            with g1:
                bolt_types = ["hex bolt", "heavy hex bolt", "allen cap", "allen csk cap",
                      "carriage bolt", "button bolt", "dome bolt", "cheese bolt"]
                bolt_type = st.selectbox("Bolt Type", bolt_types)

            with g2:
                dia_options = sorted(df_gross["DIA"].unique())
                dia = st.selectbox("Diameter (mm)", dia_options)

            with g3:
                length_options = sorted(df_gross[df_gross["DIA"] == dia]["LENGTH"].unique())
                length = st.selectbox("Length", length_options)

    # Get matching row
            row = df_gross[(df_gross["DIA"] == dia) & (df_gross["LENGTH"] == length)]

            if row.empty:
                st.error("No matching DIA + LENGTH found in JSON!")
                st.stop()

            added_mm_val = float(row["Added mm"].iloc[0])
            dia_val = float(dia)

    # L_mult based on bolt type
            formula_map = {
                "hex bolt": 3,
                "heavy hex bolt": 4,
                "allen cap": 3,
                "allen csk cap": 3.5,
                "carriage bolt": 3,
                "button bolt": 3,
                "dome bolt": 5,
                "cheese bolt": 2.5,
            }
            L_mult = formula_map.get(bolt_type, 3)

    # FINAL TOTAL FORMULA
            Total = (dia_val * dia_val * 0.0019 / 304) * (dia_val * L_mult + added_mm_val)

            st.metric("Total Weight", f"{Total:.6f}")
            # Add quantity input
            quantity_gross = st.number_input("Quantity", min_value=1, value=1, step=1, key="gross_qty")

            # Multiply Total by quantity
            Total_final = Total * quantity_gross

            st.success(f"Total × Qty = {Total_final:.6f}")


    # Mapping to description
            if "next_item" in st.session_state and pd.notna(st.session_state["next_item"]):
                selected_desc = st.session_state["next_item"]
            else:
                valid_descs = [d for d in edited_df.get("Material Description", []) if pd.notna(d) and d != ""]
                selected_desc = valid_descs[0] if valid_descs else ""

            tcol1, _ = st.columns([7,3])
            with tcol1:
                selected_desc = st.text_input("Material Description to map Total", value=selected_desc)

            if st.button("➕ Add Value (Copy Total to Rate)"):
                total_str = f"{Total_final:.6f}"
                update_rate_and_next(selected_desc, total_str)

            with st.expander("🧮 View Calculation Details"):
                st.write(f"**Diameter:** {dia_val}")
                st.write(f"**Length:** {length}")
                st.write(f"**Added mm:** {added_mm_val}")
                st.write(f"**L_mult:** {L_mult}")
                st.write(f"**Formula:** (dia² × 0.0019 / 304) × (dia×L_mult + added)")
                st.write(f"**Total = {Total:.8f}**")

# -------------------------
# TAB 6: GENERATE QUOTATION (FINAL - CLEAN + FIXED)
# -------------------------
if "df" in locals():
    with tab6:
        st.subheader("📄 Final Quotation Preview (Editable)")

        # Ensure items table exists
        if "items_df" not in st.session_state:
            st.error("❌ Rate table not available. Please complete Tab 5 first.")
            st.stop()

        # Keep memory of highlighted cells (r, col_name)
        if "highlighted_cells" not in st.session_state:
            st.session_state["highlighted_cells"] = set()

        # Base working DataFrame (copy to avoid accidental mutation)
        final_df = st.session_state["items_df"].copy()

        # Normalize column names used in UI
        rename_back = {
            "Sr. No.": "Sr No",
            "Material Description": "Description",
            "Dimension\nStandard": "Dimension Standard",
            "Material\nGrade": "Grade",
            "Qty/MOQ": "Qty"
        }
        final_df.rename(columns=rename_back, inplace=True, errors="ignore")

        # Ensure required columns exist
        for c in ["Sr No", "Item Code", "Description", "Dimension Standard", "Grade", "Finish", "Qty", "Rate"]:
            if c not in final_df.columns:
                final_df[c] = ""

        # Ensure Sr No is contiguous
        final_df["Sr No"] = list(range(1, len(final_df) + 1))

        # ---------- Highlight UI ----------
        st.markdown("### 🔴 Highlight Specific Cell (select column & row, then click)")

        # Optional small-css for making the red button compact (affects buttons globally)
        st.markdown(
            """
            <style>
            /* Make small circular-looking button for the red action (affects all st.button elements) */
            .small-red-btn button {
                min-width: 36px;
                height: 36px;
                padding: 0;
                border-radius: 18px;
                font-size: 16px;
                line-height: 0;
            }
            .small-red-btn .stButton>button:before { content: ''; }
            </style>
            """,
            unsafe_allow_html=True,
        )

        colA, colB, colC, colD = st.columns([2.4, 1.2, 1.2, 1])
        with colA:
            # Show helpful small hint (keeps layout tidy)
            st.markdown("<div style='font-size:13px;color:#444'>Choose the column and row to mark red. Repeat to mark multiple cells.</div>", unsafe_allow_html=True)

        with colB:
            highlight_col = st.selectbox(
                "Select Column",
                options=final_df.columns.tolist(),   # includes Item Code & Qty
                key="highlight_column_tab6"
            )

        with colC:
            highlight_row = st.number_input(
                "Select Row",
                min_value=1,
                max_value=max(1, len(final_df)),
                step=1,
                key="highlight_row_tab6"
            )

        with colD:
            # Put the button inside a container so CSS can target it (class applied via markdown wrapper)
            if st.button("🔴", key="make_red_tab6"):
                r = int(highlight_row) - 1
                c = highlight_col
                # store using (row_index, column_name)
                st.session_state["highlighted_cells"].add((r, c))
                st.success(f"Cell ({highlight_row}, {highlight_col}) marked RED ✔")

        # Extra controls: clear last / clear all (optional convenience)
        c1, c2 = st.columns([1, 1])
        with c1:
            if st.button("Undo Last Highlight", key="undo_last_highlight"):
                # pop last added highlight if exists (preserve order by converting to list)
                try:
                    last = list(st.session_state["highlighted_cells"])[-1]
                    st.session_state["highlighted_cells"].remove(last)
                    st.success("Removed last highlight.")
                except Exception:
                    st.info("No highlights to undo.")
        with c2:
            if st.button("Clear All Highlights", key="clear_all_highlights"):
                st.session_state["highlighted_cells"].clear()
                st.success("All highlights cleared.")

        # ---------------- MAP ITEM CODE + QTY (uses original raw inputs)
        # Use a reliable raw source (uploaded_raw_df if available) for mapping
        raw = None
        if "uploaded_raw_df" in st.session_state:
            try:
                raw = st.session_state["uploaded_raw_df"].copy()
            except Exception:
                raw = None

        if raw is None:
            try:
                raw = df.copy()
            except Exception:
                raw = None

        if raw is None:
            raw = pd.DataFrame([[""]], columns=["_"])

        # prepare normalization for raw
        cols_norm = {col: str(col).lower().replace(" ", "") for col in raw.columns}

        item_code_col = None
        qty_col = None
        desc_candidates = []

        for col, norm in cols_norm.items():
            if any(k in norm for k in ["itemcode", "itemno", "itemnumber", "code"]):
                item_code_col = col
            if any(k in norm for k in ["qty", "quantity", "moq", "set", "qty/moq"]):
                qty_col = col
            if any(k in norm for k in ["desc", "description", "material", "item", "details", "part"]):
                desc_candidates.append(col)

        if not desc_candidates:
            desc_candidates = list(raw.columns)

        raw_texts = []
        for j in range(len(raw)):
            parts = []
            for cc in desc_candidates:
                try:
                    parts.append(str(raw.at[j, cc]))
                except Exception:
                    parts.append("")
            raw_texts.append(" ".join(parts).lower())

        def extract_last_number(s):
            nums = re.findall(r"\b\d+(?:\.\d+)?\b", str(s))
            return nums[-1] if nums else ""

        try:
            pattern = item_code_pattern
        except Exception:
            pattern = re.compile(r"\b[A-Z0-9]+(?:[-/_\.]?[A-Z0-9]+){0,3}\b", re.I)

        # Run mapping on final_df (this updates Item Code and Qty when possible)
        for i in range(len(final_df)):
            desc = str(final_df.at[i, "Description"]).strip().lower() if "Description" in final_df.columns else ""
            mapped = False

            existing_code = str(final_df.at[i, "Item Code"]).strip() if "Item Code" in final_df.columns else ""
            if existing_code and item_code_col:
                for j in range(len(raw)):
                    try:
                        if str(raw.at[j, item_code_col]).strip().lower() == existing_code.lower():
                            if qty_col:
                                final_df.at[i, "Qty"] = str(raw.at[j, qty_col])
                            mapped = True
                            break
                    except Exception:
                        continue
            if mapped:
                continue

            if desc:
                tokens = [t for t in desc.split() if t]
                checks = []
                if tokens:
                    checks.append(tokens[0])
                if len(desc) >= 12:
                    checks.append(desc[:12])
                checks.append(desc)

                for j, raw_txt in enumerate(raw_texts):
                    for chk in checks:
                        if not chk:
                            continue
                        if chk in raw_txt:
                            # item code
                            if item_code_col:
                                try:
                                    val = raw.at[j, item_code_col]
                                    if pd.notna(val) and str(val).strip() != "":
                                        final_df.at[i, "Item Code"] = str(val)
                                except Exception:
                                    pass
                            else:
                                all_text = " ".join(str(raw.iloc[j].astype(str).tolist()))
                                m = pattern.search(all_text)
                                if m:
                                    final_df.at[i, "Item Code"] = m.group(0)

                            # qty
                            if qty_col:
                                try:
                                    qv = raw.at[j, qty_col]
                                    if pd.notna(qv) and str(qv).strip() != "":
                                        final_df.at[i, "Qty"] = str(qv)
                                    else:
                                        final_df.at[i, "Qty"] = extract_last_number(raw_txt)
                                except Exception:
                                    final_df.at[i, "Qty"] = extract_last_number(raw_txt)
                            else:
                                final_df.at[i, "Qty"] = extract_last_number(raw_txt)

                            mapped = True
                            break
                    if mapped:
                        break

            # fallback when lengths match
            if not mapped and len(raw) == len(final_df):
                try:
                    if item_code_col:
                        final_df.at[i, "Item Code"] = str(raw.at[i, item_code_col])
                    else:
                        all_text = " ".join(str(raw.iloc[i].astype(str).tolist()))
                        m = pattern.search(all_text)
                        if m:
                            final_df.at[i, "Item Code"] = m.group(0)

                    if qty_col:
                        final_df.at[i, "Qty"] = str(raw.at[i, qty_col])
                    else:
                        raw_txt = raw_texts[i]
                        final_df.at[i, "Qty"] = extract_last_number(raw_txt)
                except Exception:
                    pass

        # ------------------ Reorder columns AFTER mapping (important) ------------------
        desired_order = ["Sr No", "Item Code", "Description", "Dimension Standard", "Grade", "Finish", "Qty", "Rate"]
        extras = [c for c in final_df.columns if c not in desired_order]
        final_df = final_df.reindex(columns=desired_order + extras)

        # ---------- APPLY HIGHLIGHTS ----------
        df_for_display = final_df.copy()
        # Apply <font color='red'> only if not already present
        for (r, c) in list(st.session_state["highlighted_cells"]):
            try:
                # If column name no longer exists (e.g., user removed), skip
                if c not in df_for_display.columns:
                    continue
                old = str(df_for_display.at[r, c])
                if "<font color='red'>" not in old:
                    df_for_display.at[r, c] = f"<font color='red'>{old}</font>"
            except Exception:
                # If row index out of range, remove it from session to keep set clean
                try:
                    st.session_state["highlighted_cells"].remove((r, c))
                except Exception:
                    pass

        # Save for export (PDF/Excel/Word)
        st.session_state["final_table_from_tab6"] = df_for_display.copy()

        # Show editable table (user can still edit cells including colored ones)
        edited_table = st.data_editor(df_for_display, use_container_width=True, num_rows="dynamic", key="tab6_editor")
        st.session_state["final_table_edited"] = edited_table.copy()

        st.markdown("---")
        st.subheader("📥 Download")

        download_format = st.radio("Choose format to download:", ["PDF", "Excel", "Word"], horizontal=True)

        if st.button("Generate Quotation"):
            # Prefer the saved display table (with <font> tags); fallback to final_df
            export_df = st.session_state.get("final_table_from_tab6", final_df.copy())

            def _normalize_df_for_export(df):
                df2 = df.copy()
                for cc in df2.columns:
                    df2[cc] = (
                        df2[cc]
                        .astype(str)
                        .apply(
                            lambda v: html.unescape(v)
                            .replace("<span style='color:red'>", "<font color='red'>")
                            .replace("<span style=\"color:red\">", "<font color='red'>")
                            .replace("</span>", "</font>")
                        )
                    )
                return df2

            export_df = _normalize_df_for_export(export_df)

            # Build quotation dictionary (same for all formats)
            quotation = {
                "header": st.session_state["header_data"],
                "intro": "We are pleased for your esteemed inquiry. We are delighted to submit our best possible offer, compiled with our rigorous care and diligence to meet your specific requirements. We remain hopeful that our proposal aligns with your needs and look forward to your positive response.",
                "items_table": {
                    "headers": list(export_df.columns),
                    "rows": export_df.values.tolist(),
                },
                "additional_note": st.session_state.get("additional_note", ""),
                "notes": {
                    "main_note": st.session_state["quotation_notes"]["main_note"],
                    "sub_notes": st.session_state["quotation_notes"]["sub_notes"],
                    "partA": {
                        "delivery": st.session_state["quotation_notes"]["delivery"],
                        "period": st.session_state["quotation_notes"]["period"],
                        "tax": st.session_state["quotation_notes"]["tax"],
                        "pf": st.session_state["quotation_notes"]["pf"],
                        "payment": st.session_state["quotation_notes"]["payment"],
                        "validity": st.session_state["quotation_notes"]["validity"],
                    },
                    "partB": st.session_state["quotation_notes"]["bank_details"],
                    "footer": st.session_state["quotation_notes"]["footer"],
                },
                "totals": {
                    "show": st.session_state.get("show_summary_checkbox", True),
                    "subtotal": subtotal if mode == "global" else None,
                    "discount_percent": discount_percent if mode == "global" else None,
                    "discount_amount": discount_amount if mode == "global" else None,
                    "grand_total": grand_total,
                }
                }


            # ---- File Generation ----
            if download_format == "PDF":
                pdf_data = generate_pdf_dynamic(quotation)
                st.download_button("Download PDF", pdf_data, "quotation.pdf")

            elif download_format == "Excel":
                excel_data = generate_excel(quotation)
                st.download_button("Download Excel", excel_data, "quotation.xlsx")

            else:  # Word
                word_data = generate_word(quotation)
                st.download_button("Download Word", word_data, "quotation.docx")
